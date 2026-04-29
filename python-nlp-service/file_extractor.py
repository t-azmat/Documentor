"""
Centralized File Extraction Utility
Handles text extraction from DOCX, PDF, and LaTeX files
Used across all services: Grammar, Plagiarism, Citations, Formatting

Key design:
  - DOCX: reads paragraph styles (Heading 1/2/3, Title, Normal) to produce
          a structured block list instead of flat text.
  - PDF:  uses pdfplumber which exposes font-size per word; crops top/bottom
          8 % of each page to drop running headers/footers; infers headings
          from font size relative to the page's median body size.
  - TXT / LaTeX: unchanged (flat text).

The new `extract_structured()` method returns a list of blocks:
    [{"style": "Heading 1" | "Normal" | ..., "text": str}, ...]
The legacy `extract_text()` method still works — it flattens the blocks.
"""

import io
import re
import logging
import statistics
from typing import Dict, List, Optional, Tuple
import docx

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FileExtractor:
    """
    Centralized file text extraction service
    Supports: PDF, DOCX, TXT, and LaTeX files
    """
    
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.txt', '.tex', '.latex']

    # Word paragraph styles that map to heading levels
    _DOCX_HEADING_STYLES = {
        'title':     'Title',
        'heading 1': 'Heading 1',
        'heading 2': 'Heading 2',
        'heading 3': 'Heading 3',
        'heading 4': 'Heading 4',
        'heading 5': 'Heading 5',
    }
    
    def __init__(self):
        """Initialize the file extractor"""
        logger.info("File Extractor initialized")
        self._docling_converter = None
        self._docling_ready = None

    # ------------------------------------------------------------------
    # Public: structured extraction (preferred for the AI formatter)
    # ------------------------------------------------------------------

    def extract_structured(self, file_content: bytes, filename: str) -> Dict:
        """
        Extract document content as a list of typed blocks.

        Returns:
        {
          success : bool,
          blocks  : [{"style": str, "text": str}, ...],
          file_type: str,
          word_count: int,
          error   : str | None,
        }
        Block styles: "Title", "Heading 1"–"Heading 5", "Normal"
        """
        fn = filename.lower()
        try:
            extraction_backend = 'native'
            structure: List[Dict] = []
            media: List[Dict] = []
            table_of_contents: List[Dict] = []
            docling_payload = None
            if fn.endswith(('.docx', '.pdf')):
                docling_result = self._structured_from_docling(file_content, filename)
                blocks = docling_result.get('blocks', [])
                if blocks:
                    structure = docling_result.get('structure', [])
                    media = docling_result.get('media', [])
                    table_of_contents = docling_result.get('table_of_contents', [])
                    docling_payload = docling_result.get('docling')
                    extraction_backend = 'docling'
                elif fn.endswith('.docx'):
                    blocks = self._structured_from_docx(file_content)
                else:
                    blocks = self._structured_from_pdf(file_content)
            elif fn.endswith('.docx'):
                blocks = self._structured_from_docx(file_content)
            elif fn.endswith('.pdf'):
                blocks = self._structured_from_pdf(file_content)
            elif fn.endswith('.txt'):
                raw = self._decode(file_content)
                blocks = self._text_to_blocks(raw)
            elif fn.endswith(('.tex', '.latex')):
                raw = self._clean_latex(self._decode(file_content))
                blocks = self._text_to_blocks(raw)
            else:
                return {'success': False, 'blocks': [],
                        'error': f'Unsupported format', 'file_type': 'unknown'}

            word_count = sum(len(b['text'].split()) for b in blocks)
            logger.info(f"Structured extraction: {len(blocks)} blocks, {word_count} words")
            return {
                'success': True,
                'blocks': blocks,
                'structure': structure or self._blocks_to_structure(blocks),
                'media': media,
                'table_of_contents': table_of_contents or self._build_toc_from_blocks(blocks),
                'docling': docling_payload,
                'file_type': fn.rsplit('.', 1)[-1],
                'word_count': word_count,
                'backend': extraction_backend,
                'error': None,
            }
        except Exception as e:
            logger.error(f"extract_structured error: {e}")
            return {'success': False, 'blocks': [], 'error': str(e), 'file_type': 'unknown'}

    # ------------------------------------------------------------------
    # Public: legacy flat-text extraction (kept for grammar / plagiarism)
    # ------------------------------------------------------------------

    def extract_text(self, file_content: bytes, filename: str) -> Dict[str, any]:
        """
        Extract text from uploaded file with comprehensive error handling
        
        Args:
            file_content: File bytes
            filename: Name of the file
            
        Returns:
            Dictionary with:
            - success: bool
            - text: str (extracted text)
            - file_type: str (pdf/docx/txt/latex)
            - char_count: int
            - word_count: int
            - line_count: int
            - error: str (if failed)
        """
        try:
            filename_lower = filename.lower()
            file_size = len(file_content)
            logger.info(f"Extracting text from {filename} (size: {file_size} bytes)")
            
            # Determine file type and extract
            if filename_lower.endswith('.pdf'):
                result = self._extract_from_pdf(file_content, filename)
            elif filename_lower.endswith('.docx'):
                result = self._extract_from_docx(file_content, filename)
            elif filename_lower.endswith('.txt'):
                result = self._extract_from_txt(file_content, filename)
            elif filename_lower.endswith(('.tex', '.latex')):
                result = self._extract_from_latex(file_content, filename)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported file format. Supported: {", ".join(self.SUPPORTED_FORMATS)}',
                    'file_type': 'unknown'
                }
            
            if result['success']:
                # Add metadata
                text = result['text']
                result['char_count'] = len(text)
                result['word_count'] = len(text.split())
                result['line_count'] = len(text.split('\n'))
                logger.info(f"Extraction successful - {result['word_count']} words, {result['line_count']} lines")
            
            return result
            
        except Exception as e:
            logger.error(f"Unexpected error in extract_text: {str(e)}")
            return {
                'success': False,
                'error': f'Extraction failed: {str(e)}',
                'text': '',
                'file_type': 'unknown'
            }

    # ------------------------------------------------------------------
    # Structured extractors
    # ------------------------------------------------------------------

    def _get_docling_converter(self):
        if self._docling_ready is False:
            return None
        if self._docling_converter is not None:
            return self._docling_converter

        try:
            from docling.document_converter import DocumentConverter
            self._docling_converter = DocumentConverter()
            self._docling_ready = True
            logger.info("Docling converter initialized")
            return self._docling_converter
        except Exception as e:
            self._docling_ready = False
            logger.warning("Docling unavailable, falling back to native extraction: %s", e)
            return None

    def _structured_from_docling(self, file_content: bytes, filename: str) -> Dict:
        """
        Preferred extractor for PDF/DOCX using Docling.

        We intentionally normalize Docling output into the lightweight block
        schema already used by the formatter so the rest of the pipeline stays
        compatible. If Docling is unavailable or parsing fails, callers can
        fall back to the native extractor.
        """
        converter = self._get_docling_converter()
        if converter is None:
            return {}

        try:
            from docling.datamodel.base_models import InputFormat
            from docling.datamodel.document import ConversionResult
        except Exception:
            InputFormat = None
            ConversionResult = None

        stream = None
        try:
            # Newer Docling versions accept file-like streams directly.
            if InputFormat is not None:
                fmt = InputFormat.PDF if filename.lower().endswith('.pdf') else InputFormat.DOCX
                try:
                    from docling.datamodel.base_models import DocumentStream
                    stream = DocumentStream(name=filename, stream=io.BytesIO(file_content), format=fmt)
                except Exception:
                    stream = None

            result = converter.convert(stream or io.BytesIO(file_content))
            document = getattr(result, 'document', None)
            if document is None:
                return {}

            markdown = ""
            if hasattr(document, 'export_to_markdown'):
                markdown = document.export_to_markdown()
            elif hasattr(document, 'export_to_text'):
                markdown = document.export_to_text()
            elif hasattr(document, 'model_dump'):
                dumped = document.model_dump()
                markdown = dumped.get('body', '') if isinstance(dumped, dict) else ''

            markdown = (markdown or '').strip()
            if not markdown:
                return {}

            blocks = self._markdown_to_blocks(markdown)
            docling_payload = self._serialize_docling_document(document)
            structure = self._docling_to_structure(docling_payload, blocks)
            media = self._docling_to_media(docling_payload)
            toc = self._build_toc_from_blocks(blocks)
            logger.info("Docling extracted %d blocks from %s", len(blocks), filename)
            return {
                'blocks': blocks,
                'structure': structure,
                'media': media,
                'table_of_contents': toc,
                'docling': docling_payload
            }
        except Exception as e:
            logger.warning("Docling extraction failed for %s, falling back to native extractor: %s", filename, e)
            return {}

    def _structured_from_docx(self, file_content: bytes) -> List[Dict]:
        """
        Read DOCX paragraph styles to produce typed blocks.
        Heading 1/2/3 from the Word style map become structural markers.
        Tables are included as Normal blocks.
        """
        doc = docx.Document(io.BytesIO(file_content))
        blocks: List[Dict] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else 'Normal'
            mapped = self._DOCX_HEADING_STYLES.get(style_name.lower(), 'Normal')
            blocks.append({'style': mapped, 'text': text})

        # Include table cell text as Normal blocks
        for table in doc.tables:
            for row in table.rows:
                row_text = '  |  '.join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    blocks.append({'style': 'Normal', 'text': row_text})

        if not blocks:
            raise ValueError('DOCX file appears to be empty')
        return blocks

    def _structured_from_pdf(self, file_content: bytes) -> List[Dict]:
        """
        Use pdfplumber to extract text with font-size metadata.

        Strategy:
          1. Extract ALL words from every page (no blind cropping).
          2. Detect repeating headers/footers/page numbers by fingerprinting
             the top and bottom zones of each page and finding patterns that
             repeat across 40 %+ of pages.  Page 1 is handled specially since
             it often has a unique title/author block in the top zone.
          3. Detect two-column layouts (e.g. IEEE) by finding a vertical gutter
             gap in the middle 40 % of the page width.  When found, words are
             reordered: full-width content first (title/abstract area), then
             the entire left column top-to-bottom, then the entire right column
             top-to-bottom.
          4. Collect all word font-sizes across the document; compute the
             median (= body text size).
          5. A line whose average font size >= body_median * 1.18 OR whose
             fontname contains 'Bold' is treated as a heading.
          6. Lines are grouped by proximity (y-gap > 6 pt = new paragraph).
        """
        try:
            import pdfplumber
        except ImportError:
            raise ImportError(
                "pdfplumber is required for PDF extraction. "
                "Run: pip install pdfplumber"
            )

        # ==================================================================
        # Pass 1: Extract ALL words from every page — no cropping
        # ==================================================================
        raw_pages: List[List[Dict]] = []     # per-page list of word dicts
        page_heights: List[float] = []
        page_widths: List[float] = []

        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            if not pdf.pages:
                raise ValueError('PDF is empty')

            for page in pdf.pages:
                h = page.height
                pw = page.width
                page_heights.append(h)
                page_widths.append(pw)
                raw_words = page.extract_words(
                    extra_attrs=['size', 'fontname'],
                    use_text_flow=False,
                    keep_blank_chars=False,
                    x_tolerance=3,
                    y_tolerance=3,
                )
                raw_pages.append(raw_words)

        # ==================================================================
        # Pass 2: Detect repeating header/footer zones across pages
        # ==================================================================
        header_cutoffs, footer_cutoffs = self._detect_repeating_zones(
            raw_pages, page_heights
        )

        # ==================================================================
        # Pass 3: Filter words, detect columns, detect ToC, collect sizes
        # ==================================================================
        all_sizes: List[float] = []
        pages_data: List[List[Dict]] = []
        for page_idx, raw_words in enumerate(raw_pages):
            pw = page_widths[page_idx]
            h = page_heights[page_idx]
            h_cut = header_cutoffs[page_idx]
            f_cut = footer_cutoffs[page_idx]

            # Remove words in detected header/footer zones
            filtered = self._filter_header_footer_words(
                raw_words, h_cut, f_cut
            )

            # ── Column detection ────────────────────────────────────────
            gutter = self._find_column_gutter(filtered, pw)
            if gutter:
                full_w, left_w, right_w = self._split_by_column(filtered, gutter)
                ordered = full_w + left_w + right_w
                logger.debug(
                    "Two-column page %d: gutter=%.0f  full=%d  left=%d  right=%d",
                    page_idx + 1, gutter, len(full_w), len(left_w), len(right_w),
                )
            else:
                ordered = sorted(filtered,
                                 key=lambda wd: (round(wd['top'], 1), wd['x0']))

            pages_data.append(ordered)
            for wd in ordered:
                sz = wd.get('size') or 0
                if sz > 0:
                    all_sizes.append(sz)

        if not all_sizes:
            raise ValueError('PDF appears to contain no selectable text (may be scanned image)')

        body_median = statistics.median(all_sizes)
        heading_threshold = body_median * 1.18   # 18 % larger than body = heading

        # Adaptive paragraph-break gap: normal line spacing for body text
        # is roughly 1.2× font size.  A gap that is > 1.8× font size
        # (50 % bigger than normal leading) signals a true paragraph break.
        # This replaces the old hard-coded 6 pt that caused mid-sentence splits.
        para_gap_threshold = body_median * 0.5

        # ==================================================================
        # Pass 4: Group words into paragraphs, classify headings
        # ==================================================================
        blocks: List[Dict] = []
        for page_idx, words in enumerate(pages_data):
            if not words:
                continue
            # Group words into lines by their top-y coordinate (within 3 pt)
            lines: List[List[Dict]] = []
            for w in words:
                if lines and abs(w['top'] - lines[-1][0]['top']) <= 3:
                    lines[-1].append(w)
                else:
                    lines.append([w])

            # ── Adaptive paragraph grouping ─────────────────────────────
            # A paragraph break occurs when the vertical gap between the
            # bottom of the previous line and the top of the next line
            # exceeds para_gap_threshold (computed from body font size).
            paras: List[List[List[Dict]]] = []
            for line in lines:
                if paras:
                    gap = line[0]['top'] - paras[-1][-1][-1]['bottom']
                    if gap <= para_gap_threshold:
                        paras[-1].append(line)
                        continue
                paras.append([line])

            for para_lines in paras:
                flat_words = [w for line in para_lines for w in line]
                text = ' '.join(w['text'] for w in flat_words).strip()
                if not text:
                    continue

                # Determine if this paragraph is a heading
                sizes = [w.get('size') or 0 for w in flat_words if w.get('size')]
                fonts = [w.get('fontname') or '' for w in flat_words]
                avg_size = statistics.mean(sizes) if sizes else 0
                is_bold = any('bold' in f.lower() for f in fonts)
                word_count = len(text.split())

                # Check for chapter-style heading: "Chapter 1: Introduction"
                is_chapter = bool(re.match(
                    r'^chapter\s+\d+', text, re.IGNORECASE
                ))

                if is_chapter:
                    style = 'Heading 1'
                elif word_count <= 12 and (
                    avg_size >= heading_threshold or
                    (is_bold and avg_size >= body_median * 1.05)
                ):
                    # Classify heading level by relative size
                    if avg_size >= body_median * 1.45:
                        style = 'Heading 1'
                    elif avg_size >= body_median * 1.28:
                        style = 'Heading 2'
                    else:
                        style = 'Heading 3'
                else:
                    style = 'Normal'

                # Fix hyphenated line-breaks (e.g. "meth-\nodology" → "methodology")
                text = re.sub(r'(\w)-\s+(\w)', r'\1\2', text)
                blocks.append({'style': style, 'text': text})

        if not blocks:
            raise ValueError('No text could be extracted from the PDF')
        return blocks

    def _markdown_to_blocks(self, markdown: str) -> List[Dict]:
        """
        Convert structure-rich markdown into formatter blocks.

        This keeps paragraph boundaries, headings, lists, and table text
        instead of collapsing everything into one plain-text blob.
        """
        blocks: List[Dict] = []
        paragraph_lines: List[str] = []
        table_lines: List[str] = []
        in_code_block = False

        def flush_paragraph():
            nonlocal paragraph_lines
            text = '\n'.join(line.rstrip() for line in paragraph_lines).strip()
            if text:
                blocks.append({'style': 'Normal', 'text': text})
            paragraph_lines = []

        def flush_table():
            nonlocal table_lines
            text = '\n'.join(line.rstrip() for line in table_lines).strip()
            if text:
                blocks.append({'style': 'Normal', 'text': text})
            table_lines = []

        for raw_line in markdown.splitlines():
            line = raw_line.rstrip()
            stripped = line.strip()

            if stripped.startswith('```'):
                in_code_block = not in_code_block
                if not in_code_block:
                    flush_paragraph()
                continue

            if in_code_block:
                paragraph_lines.append(line)
                continue

            if not stripped:
                flush_paragraph()
                flush_table()
                continue

            heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if heading_match:
                flush_paragraph()
                flush_table()
                level = len(heading_match.group(1))
                blocks.append({
                    'style': f'Heading {min(level, 5)}',
                    'text': heading_match.group(2).strip(),
                    'section_level': level
                })
                continue

            if '|' in stripped and len(stripped.split('|')) >= 3:
                flush_paragraph()
                table_lines.append(stripped)
                continue

            bullet_match = re.match(r'^[-*+]\s+(.+)$', stripped)
            if bullet_match:
                flush_table()
                blocks.append({'style': 'Normal', 'text': f"• {bullet_match.group(1).strip()}"})
                continue

            numbered_match = re.match(r'^\d+[\.\)]\s+(.+)$', stripped)
            if numbered_match:
                flush_table()
                blocks.append({'style': 'Normal', 'text': stripped})
                continue

            paragraph_lines.append(stripped)

        flush_paragraph()
        flush_table()
        return blocks

    def _build_toc_from_blocks(self, blocks: List[Dict]) -> List[Dict]:
        toc = []
        for idx, block in enumerate(blocks):
            text = str(block.get('text', '')).strip()
            if not text:
                continue
            level = block.get('section_level', 0)
            if not level:
                style = str(block.get('style', ''))
                match = re.search(r'Heading\s*(\d+)', style, re.IGNORECASE)
                if match:
                    level = int(match.group(1))
                elif style == 'Title':
                    level = 1
            if level:
                toc.append({
                    'level': int(level),
                    'text': text[:200],
                    'index': idx
                })
        return toc

    def _blocks_to_structure(self, blocks: List[Dict]) -> List[Dict]:
        structure = []
        for block in blocks:
            text = str(block.get('text', '')).strip()
            if not text:
                continue
            style = str(block.get('style', 'Normal'))
            section_level = block.get('section_level', 0)

            heading_match = re.search(r'Heading\s*(\d+)', style, re.IGNORECASE)
            if style == 'Title':
                structure.append({'type': 'heading1', 'content': text, 'level': 1})
                continue
            if heading_match or section_level:
                level = int(section_level or heading_match.group(1))
                structure.append({
                    'type': f'heading{min(max(level, 1), 3)}',
                    'content': text,
                    'level': level
                })
                continue
            if text.startswith('â€¢ ') or text.startswith('• '):
                bullet_text = text[2:].strip()
                structure.append({'type': 'list-item', 'content': bullet_text})
                continue
            if text.startswith('|') and text.count('|') >= 2:
                rows = self._parse_markdown_table(text)
                if rows:
                    structure.append({
                        'type': 'table',
                        'content': rows,
                        'tableData': {
                            'rows': len(rows),
                            'columns': max((len(r) for r in rows), default=0),
                            'data': rows
                        }
                    })
                    continue
            structure.append({'type': 'paragraph', 'content': text})
        return structure

    def _parse_markdown_table(self, text: str) -> List[List[str]]:
        rows = []
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped.startswith('|'):
                continue
            if re.match(r'^\|?[\s:\-]+\|[\s|:\-]*$', stripped):
                continue
            cells = [cell.strip() for cell in stripped.strip('|').split('|')]
            if any(cells):
                rows.append(cells)
        return rows

    def _serialize_docling_document(self, document) -> Optional[Dict]:
        try:
            if hasattr(document, 'model_dump'):
                dumped = document.model_dump()
                if isinstance(dumped, dict):
                    return dumped
            if hasattr(document, 'dict'):
                dumped = document.dict()
                if isinstance(dumped, dict):
                    return dumped
            if hasattr(document, '__dict__'):
                dumped = {
                    key: value for key, value in document.__dict__.items()
                    if not key.startswith('_')
                }
                if dumped:
                    return dumped
        except Exception as e:
            logger.debug("Failed to serialize Docling document: %s", e)
        return None

    def _docling_to_structure(self, docling_payload: Optional[Dict], fallback_blocks: List[Dict]) -> List[Dict]:
        if not isinstance(docling_payload, dict):
            return self._blocks_to_structure(fallback_blocks)

        structure = []

        def visit(node):
            if isinstance(node, dict):
                label = str(node.get('label') or node.get('type') or '').lower()
                text = str(
                    node.get('text')
                    or node.get('orig')
                    or node.get('content')
                    or node.get('name')
                    or ''
                ).strip()

                if label in {'section_header', 'heading', 'title'} and text:
                    level = int(node.get('level') or 1)
                    structure.append({
                        'type': f'heading{min(max(level, 1), 3)}',
                        'content': text,
                        'level': level
                    })
                elif label in {'text', 'paragraph', 'caption', 'list_item'} and text:
                    structure.append({
                        'type': 'list-item' if label == 'list_item' else 'paragraph',
                        'content': text
                    })
                elif 'table' in label:
                    rows = self._extract_rows_from_node(node)
                    if rows:
                        structure.append({
                            'type': 'table',
                            'content': rows,
                            'tableData': {
                                'rows': len(rows),
                                'columns': max((len(r) for r in rows), default=0),
                                'data': rows
                            }
                        })

                for value in node.values():
                    visit(value)
            elif isinstance(node, list):
                for item in node:
                    visit(item)

        visit(docling_payload)
        return structure or self._blocks_to_structure(fallback_blocks)

    def _extract_rows_from_node(self, node: Dict) -> List[List[str]]:
        rows = []
        data = node.get('data') or node.get('table_data') or node.get('table_cells') or node.get('cells')
        if isinstance(data, list):
            for row in data:
                if isinstance(row, dict):
                    cells = row.get('cells') or row.get('data') or []
                    if isinstance(cells, list):
                        rows.append([str(cell.get('text') if isinstance(cell, dict) else cell).strip() for cell in cells])
                elif isinstance(row, list):
                    rows.append([str(cell.get('text') if isinstance(cell, dict) else cell).strip() for cell in row])
        return [row for row in rows if any(cell for cell in row)]

    def _docling_to_media(self, docling_payload: Optional[Dict]) -> List[Dict]:
        if not isinstance(docling_payload, dict):
            return []

        media = []

        def visit(node):
            if isinstance(node, dict):
                label = str(node.get('label') or node.get('type') or '').lower()
                if any(keyword in label for keyword in ('picture', 'image', 'figure')):
                    caption = str(node.get('caption') or node.get('text') or node.get('name') or '').strip()
                    media.append({
                        'type': 'image',
                        'filename': node.get('filename') or '',
                        'caption': caption,
                        'altText': caption,
                        'relativePath': ''
                    })
                for value in node.values():
                    visit(value)
            elif isinstance(node, list):
                for item in node:
                    visit(item)

        visit(docling_payload)
        return media[:50]

    # ------------------------------------------------------------------
    # Table of Contents detection
    # ------------------------------------------------------------------

    @staticmethod
    def _is_toc_line(text: str) -> bool:
        """
        Return True if a text string looks like a Table of Contents entry.

        Typical ToC line patterns (after PDF word extraction):
          - "Introduction ........... 8"  → text + dot leaders + page number
          - "1.1 Background 12"          → numbered heading + page number
          - "Chapter 1: Introduction 5"  → chapter prefix + page number
          - "Table of Contents"          → the ToC heading itself
          - "8 Problem Statement"        → page number + heading (reversed by PDF)
          - Standalone page numbers      → "3", "12", "iv"

        The key signal: short text (≤ 15 words) ending with a bare number
        (the page number) after stripping dot leaders.
        """
        if not text or not text.strip():
            return False
        t = text.strip()

        # The ToC heading itself
        if re.match(r'^(?:table\s+of\s+contents|contents|list\s+of\s+(?:figures|tables))$',
                     t, re.IGNORECASE):
            return True

        # Standalone page number (possibly Roman numerals)
        if re.match(r'^(?:\d{1,4}|[ivxlc]{1,6})$', t, re.IGNORECASE):
            return True

        # Strip dot leaders / underscores that act as visual fill
        cleaned = re.sub(r'[._·…\u2026]{3,}', ' ', t)
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()

        words = cleaned.split()
        if not words:
            return False

        # A ToC entry typically ends with a page number
        last = words[-1]
        ends_with_number = bool(re.match(r'^\d{1,4}$', last))

        # Also handle entries where the page number comes first:
        # e.g. "8 Problem Statement" (PDF extraction order quirk)
        first = words[0]
        starts_with_number = bool(re.match(r'^\d{1,4}$', first)) and len(words) >= 2

        if not ends_with_number and not starts_with_number:
            return False

        # ToC entries are short — typically ≤ 12 words of actual content
        content_words = len(words) - 1  # minus the page number
        if content_words > 15:
            return False

        # If it looks like a real sentence (ends with period before the number),
        # it's probably body text, not a ToC entry.
        content_text = ' '.join(words[:-1]) if ends_with_number else ' '.join(words[1:])
        if content_text and content_text[-1] in '.!?':
            return False

        return True

    @staticmethod
    def _is_toc_page(words: List[Dict], page_width: float) -> bool:
        """
        Return True if a page full of words looks like a Table of Contents.

        A ToC page has these telltale characteristics:
          1. Many short "lines" (the heading text) on the left half of the page
             paired with isolated page numbers on the right half.
          2. A large fraction of lines match the _is_toc_line pattern.
          3. The heading "Table of Contents" / "Contents" appears at the top.
          4. High density of isolated small numbers (page references).

        The heuristic: group words into lines (by y-coordinate), build each
        line's text, and check what fraction look like ToC entries.
        If >= 50 % of lines with any real text are ToC-like, declare the page
        as a ToC page.
        """
        if len(words) < 10:
            return False

        # Group into lines by y-coordinate (within 3 pt)
        lines_of_words: List[List[Dict]] = []
        sorted_words = sorted(words, key=lambda w: (round(w['top'], 1), w['x0']))
        for w in sorted_words:
            if lines_of_words and abs(w['top'] - lines_of_words[-1][0]['top']) <= 3:
                lines_of_words[-1].append(w)
            else:
                lines_of_words.append([w])

        if not lines_of_words:
            return False

        # Check for explicit "Table of Contents" heading in the first few lines
        has_toc_heading = False
        for line_words in lines_of_words[:5]:
            line_text = ' '.join(w['text'] for w in line_words).strip()
            if re.match(r'^(?:table\s+of\s+contents|contents)$', line_text, re.IGNORECASE):
                has_toc_heading = True
                break

        # Build line texts and count ToC-like lines
        toc_count = 0
        real_text_lines = 0
        isolated_numbers = 0

        for line_words in lines_of_words:
            line_text = ' '.join(w['text'] for w in line_words).strip()
            if not line_text:
                continue

            # Count standalone numbers (page references scattered across the page)
            if re.match(r'^\d{1,4}$', line_text):
                isolated_numbers += 1
                toc_count += 1
                real_text_lines += 1
                continue

            real_text_lines += 1
            if FileExtractor._is_toc_line(line_text):
                toc_count += 1

        if real_text_lines < 5:
            return False

        toc_ratio = toc_count / real_text_lines

        # With explicit "Table of Contents" heading: 40 % threshold
        # Without explicit heading: 55 % threshold (stricter)
        threshold = 0.40 if has_toc_heading else 0.55

        is_toc = toc_ratio >= threshold

        if is_toc:
            logger.debug(
                "ToC page: %d/%d lines match (%.0f%%), heading=%s",
                toc_count, real_text_lines, toc_ratio * 100, has_toc_heading,
            )

        return is_toc

    # ------------------------------------------------------------------
    # Header / footer detection via cross-page repetition
    # ------------------------------------------------------------------

    # Zone height as fraction of page height to examine for headers/footers
    _ZONE_FRAC = 0.12          # top/bottom 12 %
    # Minimum fraction of pages a zone fingerprint must appear on
    _REPEAT_THRESHOLD = 0.40   # 40 % of pages
    # Fallback crop when document is a single page (can't do repetition)
    _SINGLE_PAGE_CROP = 0.06   # 6 % top/bottom

    @staticmethod
    def _build_zone_fingerprint(words: List[Dict]) -> str:
        """
        Build a normalised fingerprint from a list of words.
        - Replaces all digit sequences with '#' so page numbers don't
          break matching (e.g. "Vol 12 Page 3" and "Vol 12 Page 7"
          both become "vol ## page #").
        - Lowercases, collapses whitespace.
        """
        text = ' '.join(w.get('text', '') for w in words).strip()
        if not text:
            return ''
        text = re.sub(r'\d+', '#', text)
        text = re.sub(r'\s+', ' ', text).strip().lower()
        return text

    def _detect_repeating_zones(
        self,
        raw_pages: List[List[Dict]],
        page_heights: List[float],
    ) -> Tuple[List[float], List[float]]:
        """
        Detect repeating header and footer zones across all pages.

        Returns two lists of the same length as raw_pages:
          header_cutoffs[i] — y-coordinate below which header words should be
                              excluded on page i  (0.0 = no header detected)
          footer_cutoffs[i] — y-coordinate above which footer words should be
                              excluded on page i  (page_height = no footer)

        Algorithm:
          1. For each page, fingerprint the words in the top 12 % and bottom
             12 % zones (digits normalised to '#').
          2. Count how many pages share the same fingerprint.
          3. If a fingerprint appears on >= 40 % of pages, mark those pages
             as having a header/footer.
          4. For marked pages, compute the actual y-boundary from the word
             positions (max bottom of header words / min top of footer words)
             plus a small padding.
          5. Page 1 special handling: the top zone is only excluded if its
             fingerprint also appears on page 2 (protects title blocks).
        """
        n_pages = len(raw_pages)
        if n_pages == 0:
            return [], []

        # Default: no exclusion
        header_cutoffs = [0.0] * n_pages
        footer_cutoffs = list(page_heights)

        # Single-page document: can't do repetition detection, use a
        # conservative fixed crop as fallback.
        if n_pages == 1:
            h = page_heights[0]
            header_cutoffs[0] = h * self._SINGLE_PAGE_CROP
            footer_cutoffs[0] = h * (1.0 - self._SINGLE_PAGE_CROP)
            logger.debug("Single page — using %.0f%% fixed crop", self._SINGLE_PAGE_CROP * 100)
            return header_cutoffs, footer_cutoffs

        # Build per-page zone fingerprints
        top_fps: List[str] = []
        bot_fps: List[str] = []
        top_zone_words: List[List[Dict]] = []
        bot_zone_words: List[List[Dict]] = []

        for page_idx, words in enumerate(raw_pages):
            h = page_heights[page_idx]
            top_boundary = h * self._ZONE_FRAC
            bot_boundary = h * (1.0 - self._ZONE_FRAC)

            tw = [w for w in words if w.get('top', 0) < top_boundary]
            bw = [w for w in words if w.get('bottom', h) > bot_boundary]

            top_zone_words.append(tw)
            bot_zone_words.append(bw)
            top_fps.append(self._build_zone_fingerprint(tw))
            bot_fps.append(self._build_zone_fingerprint(bw))

        # Count fingerprint occurrences (skip empty fingerprints)
        from collections import Counter
        min_repeats = max(2, int(n_pages * self._REPEAT_THRESHOLD))

        top_counts = Counter(fp for fp in top_fps if fp)
        bot_counts = Counter(fp for fp in bot_fps if fp)

        repeating_top = {fp for fp, cnt in top_counts.items() if cnt >= min_repeats}
        repeating_bot = {fp for fp, cnt in bot_counts.items() if cnt >= min_repeats}

        logger.info(
            "Header/footer detection: %d pages, %d repeating top zones, "
            "%d repeating bottom zones",
            n_pages, len(repeating_top), len(repeating_bot),
        )

        # Compute per-page cutoffs
        _PAD = 3.0  # pt padding below header / above footer boundary

        for page_idx in range(n_pages):
            h = page_heights[page_idx]

            # ── Header ──────────────────────────────────────────────────
            if top_fps[page_idx] in repeating_top:
                # Page 1 special case: only exclude if page 2 has the same
                # fingerprint — otherwise the top zone is likely the title.
                if page_idx == 0 and n_pages > 1:
                    if top_fps[0] != top_fps[1]:
                        # Title / author block — don't exclude
                        continue

                tw = top_zone_words[page_idx]
                if tw:
                    max_bottom = max(w.get('bottom', 0) for w in tw)
                    header_cutoffs[page_idx] = max_bottom + _PAD
                    logger.debug(
                        "Page %d: header cutoff at y=%.1f (fingerprint: %s…)",
                        page_idx + 1, header_cutoffs[page_idx],
                        top_fps[page_idx][:40],
                    )

            # ── Footer ──────────────────────────────────────────────────
            if bot_fps[page_idx] in repeating_bot:
                bw = bot_zone_words[page_idx]
                if bw:
                    min_top = min(w.get('top', h) for w in bw)
                    footer_cutoffs[page_idx] = min_top - _PAD
                    logger.debug(
                        "Page %d: footer cutoff at y=%.1f (fingerprint: %s…)",
                        page_idx + 1, footer_cutoffs[page_idx],
                        bot_fps[page_idx][:40],
                    )

        return header_cutoffs, footer_cutoffs

    @staticmethod
    def _filter_header_footer_words(
        words: List[Dict],
        header_cutoff: float,
        footer_cutoff: float,
    ) -> List[Dict]:
        """
        Remove words that fall within detected header or footer zones.
        A word is excluded if its vertical midpoint is above the header
        cutoff or below the footer cutoff.
        """
        if header_cutoff <= 0 and footer_cutoff >= 9999:
            return words
        filtered = []
        for w in words:
            y_mid = (w.get('top', 0) + w.get('bottom', 0)) / 2.0
            if y_mid < header_cutoff:
                continue
            if y_mid > footer_cutoff:
                continue
            filtered.append(w)
        return filtered

    # ------------------------------------------------------------------
    # Helper: two-column layout detection
    # ------------------------------------------------------------------

    @staticmethod
    def _find_column_gutter(words: List[Dict], page_width: float) -> Optional[float]:
        """
        Detect the x-coordinate of a two-column gutter, if one exists.

        Algorithm:
          - Focus on word x-midpoints in the central 40 % of the page (x in
            30 %–70 % of page width) — that is where a gutter must lie.
          - Find the largest gap between consecutive x-midpoints in that zone.
          - Accept it as a gutter only when it is >= 10 pt wide  (≈ 1/7 inch,
            which is narrower than any real IEEE gutter of ~17 pt).

        Returns gutter centre x, or None if the page appears single-column.
        """
        if len(words) < 20:
            return None

        lo, hi = page_width * 0.30, page_width * 0.70
        central = sorted(
            (w['x0'] + w['x1']) / 2.0
            for w in words
            if lo <= (w['x0'] + w['x1']) / 2.0 <= hi
        )

        if len(central) < 6:
            # Very sparse centre — likely already a wide gutter; fall back to
            # checking whether words cluster on both sides of the midpoint.
            left_n  = sum(1 for w in words if (w['x0'] + w['x1']) / 2.0 < page_width * 0.50)
            right_n = sum(1 for w in words if (w['x0'] + w['x1']) / 2.0 > page_width * 0.50)
            if left_n >= 10 and right_n >= 10:
                return page_width * 0.50
            return None

        gaps = [
            (central[i + 1] - central[i], (central[i] + central[i + 1]) / 2.0)
            for i in range(len(central) - 1)
        ]
        max_gap_size, max_gap_centre = max(gaps, key=lambda g: g[0])

        if max_gap_size < 10:
            return None

        logger.debug(
            "Two-column gutter at x=%.1f  gap=%.1f pt  page_w=%.0f",
            max_gap_centre, max_gap_size, page_width,
        )
        return max_gap_centre

    @staticmethod
    def _split_by_column(
        words: List[Dict], gutter: float
    ) -> Tuple[List[Dict], List[Dict], List[Dict]]:
        """
        Partition words into three ordered lists:
          full_width  — words that physically straddle the gutter
                        (title block, abstract box, wide figures/tables)
          left_col    — words entirely left of the gutter
          right_col   — words entirely right of the gutter

        Each list is sorted top→bottom then left→right so that subsequent
        line-grouping logic produces correct reading order within each region.
        """
        key = lambda w: (round(w['top'], 1), w['x0'])
        full_width, left_col, right_col = [], [], []
        for w in words:
            x_mid = (w['x0'] + w['x1']) / 2.0
            if w['x0'] < gutter and w['x1'] > gutter:
                full_width.append(w)
            elif x_mid <= gutter:
                left_col.append(w)
            else:
                right_col.append(w)
        return sorted(full_width, key=key), sorted(left_col, key=key), sorted(right_col, key=key)

    # ------------------------------------------------------------------
    # Helper: convert flat text to blocks (used for TXT / LaTeX)
    # ------------------------------------------------------------------

    @staticmethod
    def _text_to_blocks(text: str) -> List[Dict]:
        """
        Heuristically annotate plain text with heading/normal styles.
        Rules (in priority order):
          0. "Chapter N: Title"  → Heading 1
          1. ALL-CAPS short line → Heading 1
          2. Numbered heading    → Heading 2  (e.g. "1. Introduction")
          3. Short line followed by blank line → Heading 3
          4. Everything else     → Normal
        """
        lines = text.splitlines()
        blocks: List[Dict] = []
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            wc = len(line.split())

            # Chapter heading: "Chapter 1: Introduction" / "Chapter 2 Background"
            if re.match(r'^chapter\s+\d+', line, re.IGNORECASE) and wc <= 12:
                blocks.append({'style': 'Heading 1', 'text': line})
            # ALL-CAPS heading
            elif wc <= 8 and line == line.upper() and re.search(r'[A-Z]', line):
                blocks.append({'style': 'Heading 1', 'text': line.title()})
            # Numbered section like "1. Introduction" or "2.1 Related Work"
            elif re.match(r'^\d+(\.\d+)?\s+[A-Z]', line) and wc <= 10:
                blocks.append({'style': 'Heading 2', 'text': re.sub(r'^\d+(\.\d+)?\s+', '', line)})
            # Short line followed by blank — likely a heading
            elif wc <= 6 and i + 1 < len(lines) and not lines[i + 1].strip():
                blocks.append({'style': 'Heading 3', 'text': line})
            else:
                blocks.append({'style': 'Normal', 'text': line})
            i += 1
        return blocks

    @staticmethod
    def _decode(file_content: bytes) -> str:
        for enc in ('utf-8', 'latin-1', 'windows-1252'):
            try:
                return file_content.decode(enc)
            except UnicodeDecodeError:
                continue
        raise ValueError('Could not decode file content')
    
    def _extract_from_pdf(self, file_content: bytes, filename: str) -> Dict:
        """Extract flat text from PDF using pdfplumber (legacy path for grammar/plagiarism)."""
        try:
            blocks = self._structured_from_pdf(file_content)
            text = '\n'.join(
                ('\n' + b['text'] if b['style'].startswith('Heading') else b['text'])
                for b in blocks
            )
            if not text.strip():
                return {'success': False, 'error': 'PDF appears to be empty or image-only',
                        'text': '', 'file_type': 'pdf'}
            logger.info(f"PDF extraction complete - {len(text)} characters")
            return {'success': True, 'text': text, 'file_type': 'pdf'}
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            return {'success': False, 'error': f'Error reading PDF: {str(e)}',
                    'text': '', 'file_type': 'pdf'}
    
    def _extract_from_docx(self, file_content: bytes, filename: str) -> Dict:
        """Extract flat text from DOCX (legacy path for grammar/plagiarism)."""
        try:
            blocks = self._structured_from_docx(file_content)
            text = '\n'.join(
                ('\n' + b['text'] if b['style'].startswith('Heading') or b['style'] == 'Title' else b['text'])
                for b in blocks
            )
            if not text.strip():
                return {'success': False, 'error': 'DOCX file appears to be empty',
                        'text': '', 'file_type': 'docx'}
            logger.info(f"DOCX extraction complete - {len(text)} characters")
            return {'success': True, 'text': text, 'file_type': 'docx'}
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            return {'success': False, 'error': f'Error reading DOCX: {str(e)}',
                    'text': '', 'file_type': 'docx'}
    
    def _extract_from_txt(self, file_content: bytes, filename: str) -> Dict:
        """Extract text from TXT file"""
        try:
            text = self._decode(file_content)
            
            if not text.strip():
                return {
                    'success': False,
                    'error': 'Text file is empty',
                    'text': '',
                    'file_type': 'txt'
                }
            
            logger.info(f"TXT extraction complete - {len(text)} characters")
            return {
                'success': True,
                'text': text,
                'file_type': 'txt'
            }
            
        except Exception as e:
            logger.error(f"TXT decoding error: {str(e)}")
            return {
                'success': False,
                'error': f'Error reading TXT file: {str(e)}',
                'text': '',
                'file_type': 'txt'
            }
    
    def _extract_from_latex(self, file_content: bytes, filename: str) -> Dict:
        """Extract text from LaTeX file, removing LaTeX commands"""
        try:
            # Decode LaTeX file
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                text = file_content.decode('latin-1')
            
            # Remove LaTeX commands but keep content
            cleaned_text = self._clean_latex(text)
            
            if not cleaned_text.strip():
                return {
                    'success': False,
                    'error': 'LaTeX file appears to be empty after processing',
                    'text': '',
                    'file_type': 'latex'
                }
            
            logger.info(f"LaTeX extraction complete - {len(cleaned_text)} characters")
            return {
                'success': True,
                'text': cleaned_text,
                'file_type': 'latex',
                'original_latex': text  # Keep original for reference
            }
            
        except Exception as e:
            logger.error(f"LaTeX extraction error: {str(e)}")
            return {
                'success': False,
                'error': f'Error reading LaTeX file: {str(e)}',
                'text': '',
                'file_type': 'latex'
            }
    
    def _clean_latex(self, latex_text: str) -> str:
        """
        Remove LaTeX commands and keep readable text
        
        Args:
            latex_text: Raw LaTeX text
            
        Returns:
            Cleaned text without LaTeX commands
        """
        text = latex_text
        
        # Remove comments
        text = re.sub(r'%.*?\n', '\n', text)
        
        # Remove preamble (everything before \begin{document})
        doc_start = text.find(r'\begin{document}')
        if doc_start != -1:
            text = text[doc_start:]
        
        # Remove \end{document} and everything after
        doc_end = text.find(r'\end{document}')
        if doc_end != -1:
            text = text[:doc_end]
        
        # Remove common commands but keep their content
        # \textbf{text} -> text
        text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\emph\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\underline\{([^}]*)\}', r'\1', text)
        
        # Remove section commands but keep titles
        text = re.sub(r'\\section\{([^}]*)\}', r'\n\n\1\n', text)
        text = re.sub(r'\\subsection\{([^}]*)\}', r'\n\1\n', text)
        text = re.sub(r'\\subsubsection\{([^}]*)\}', r'\n\1\n', text)
        text = re.sub(r'\\chapter\{([^}]*)\}', r'\n\n\1\n\n', text)
        
        # Remove citations but keep them marked
        text = re.sub(r'\\cite\{([^}]*)\}', r'[Citation: \1]', text)
        text = re.sub(r'\\citep\{([^}]*)\}', r'[\1]', text)
        text = re.sub(r'\\citet\{([^}]*)\}', r'\1', text)
        
        # Remove labels and refs
        text = re.sub(r'\\label\{[^}]*\}', '', text)
        text = re.sub(r'\\ref\{([^}]*)\}', r'[ref]', text)
        
        # Remove begin/end environments but keep content
        text = re.sub(r'\\begin\{([^}]*)\}', '', text)
        text = re.sub(r'\\end\{([^}]*)\}', '', text)
        
        # Remove remaining simple commands
        text = re.sub(r'\\[a-zA-Z]+(\[[^\]]*\])?\{([^}]*)\}', r'\2', text)
        
        # Remove remaining backslash commands
        text = re.sub(r'\\[a-zA-Z]+', '', text)
        
        # Clean up extra whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        return text.strip()
    
    def is_supported(self, filename: str) -> bool:
        """
        Check if file format is supported
        
        Args:
            filename: Name of the file
            
        Returns:
            True if supported, False otherwise
        """
        filename_lower = filename.lower()
        return any(filename_lower.endswith(fmt) for fmt in self.SUPPORTED_FORMATS)
    
    def get_file_type(self, filename: str) -> Optional[str]:
        """
        Get file type from filename
        
        Args:
            filename: Name of the file
            
        Returns:
            File type (pdf/docx/txt/latex) or None
        """
        filename_lower = filename.lower()
        if filename_lower.endswith('.pdf'):
            return 'pdf'
        elif filename_lower.endswith('.docx'):
            return 'docx'
        elif filename_lower.endswith('.txt'):
            return 'txt'
        elif filename_lower.endswith(('.tex', '.latex')):
            return 'latex'
        return None
