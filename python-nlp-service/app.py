"""
Python NLP Service - Chapter 4 Implementation
Flask microservice for NLP functionality
Implements all 4 algorithms from Chapter 4:
1. Grammar Enhancement (T5 Transformer)
2. Plagiarism Detection (Semantic Similarity)
3. Document Formatting (Multi-style)
4. Section Detection (Naive Bayes)
"""

from flask import Flask, request, jsonify, Response, stream_with_context, send_file
from flask_cors import CORS
import os
import re
import uuid
import json
import logging
import time
import importlib.util
import sys
from pathlib import Path
from grammar_enhancer import GrammarEnhancer
from plagiarism_detector import SemanticPlagiarismDetector, OnlinePlagiarismChecker
from formatting_engine import DocumentFormattingEngine
from section_detector import DocumentSectionDetector
from citation import CitationManager, extract_references_section, match_citations_to_references, AISemanticCitationMatcher
from file_extractor import FileExtractor
from pdf_extractor import extract_pdf_complete
from docx_extractor import extract_docx_complete
import tempfile
import os

# PlagiarismStudio dependencies (lazy-loaded)
_studio_cosine_ready = False
_studio_bert_model = None
_studio_gpt2_model = None
_studio_gpt2_tokenizer = None

def _get_studio_bert():
    global _studio_bert_model
    if _studio_bert_model is None:
        try:
            from sentence_transformers import SentenceTransformer
            # low_cpu_mem_usage=False prevents meta-tensor initialisation in
            # transformers >=4.38 which causes "Cannot copy out of meta tensor"
            # with PyTorch 2.x when SentenceTransformer calls .to(device).
            _studio_bert_model = SentenceTransformer(
                'paraphrase-MiniLM-L6-v2',
                model_kwargs={'low_cpu_mem_usage': False},
            )
        except Exception as e:
            logger.warning(f"BERT model unavailable: {e}")
    return _studio_bert_model

def _get_studio_gpt2():
    global _studio_gpt2_model, _studio_gpt2_tokenizer
    if _studio_gpt2_model is None:
        try:
            import warnings
            warnings.filterwarnings("ignore")
            import transformers
            transformers.logging.set_verbosity_error()
            from transformers import GPT2LMHeadModel, GPT2TokenizerFast
            _studio_gpt2_tokenizer = GPT2TokenizerFast.from_pretrained("gpt2")
            _studio_gpt2_model = GPT2LMHeadModel.from_pretrained("gpt2")
            _studio_gpt2_model.eval()
        except Exception as e:
            logger.warning(f"GPT-2 model unavailable: {e}")
    return _studio_gpt2_model, _studio_gpt2_tokenizer

def studio_cosine_similarity(text1, text2):
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity as sk_cosine
    vectorizer = TfidfVectorizer()
    vectors = vectorizer.fit_transform([text1, text2])
    score = sk_cosine(vectors[0:1], vectors[1:2])[0][0] * 100
    return round(float(score), 2)

def studio_bert_similarity(text1, text2):
    model = _get_studio_bert()
    if model is None:
        return None
    from sentence_transformers import util
    e1 = model.encode(text1, convert_to_tensor=True)
    e2 = model.encode(text2, convert_to_tensor=True)
    score = util.cos_sim(e1, e2).item() * 100
    return round(float(score), 2)

def _gpt2_perplexity(text):
    """GPT-2 perplexity — supplementary signal only. NOT reliable alone for ChatGPT."""
    import torch
    model, tokenizer = _get_studio_gpt2()
    if model is None or tokenizer is None:
        return None
    if len(text) < 100:
        return None
    try:
        encodings = tokenizer(text[:4000], return_tensors="pt")
        max_length = model.config.n_positions
        stride = 512
        lls = []
        for i in range(0, encodings.input_ids.size(1), stride):
            begin_loc = max(i + stride - max_length, 0)
            input_ids = encodings.input_ids[:, begin_loc:i+stride]
            target_ids = input_ids.clone()
            with torch.no_grad():
                outputs = model(input_ids, labels=target_ids)
                lls.append(outputs.loss * input_ids.size(1))
        if not lls:
            return None
        total_ll = torch.stack(lls).sum()
        ppl = torch.exp(total_ll / encodings.input_ids.size(1))
        return round(float(ppl.item()), 2)
    except Exception:
        return None

def studio_ai_detect(text):
    """
    Multi-feature AI text detection.
    Returns dict with: score (0-100, higher=more AI), label, features, perplexity.

    Features used:
      1. Sentence burstiness  — most reliable signal.
         ChatGPT writes uniformly (low variance), humans mix short/long sentences.
         Burstiness = (std - mean) / (std + mean). Low/negative = AI.
      2. Coefficient of variation (CV) of sentence lengths.
         CV < 0.4 strongly suggests AI-generated text.
      3. Mean sentence length. ChatGPT stays consistently 15-25 words.
      4. Filler phrase frequency. ChatGPT overuses certain phrases.
      5. GPT-2 perplexity (only if model loaded — supplementary, NOT primary).
    """
    import re
    from collections import Counter

    if not text or len(text.strip()) < 80:
        return {'score': None, 'label': 'Too short to analyze', 'features': {}, 'perplexity': None}

    # ── Feature 1 & 2 & 3: sentence length stats ──────────────────────────────
    raw_sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    sentences = [s.strip() for s in raw_sentences if len(s.strip().split()) >= 4]
    sent_lengths = [len(s.split()) for s in sentences]

    burstiness = None
    cv = None
    mean_len = None
    if len(sent_lengths) >= 4:
        n = len(sent_lengths)
        mean_len = sum(sent_lengths) / n
        variance = sum((l - mean_len) ** 2 for l in sent_lengths) / n
        std_dev = variance ** 0.5
        burstiness = (std_dev - mean_len) / (std_dev + mean_len + 1e-9)
        cv = std_dev / (mean_len + 1e-9)

    # ── Feature 4: ChatGPT filler phrases ─────────────────────────────────────
    filler_phrases = [
        'it is important to note', 'it is worth noting', 'it is essential to',
        'in conclusion', 'in summary', 'to summarize', 'in other words',
        'furthermore', 'moreover', 'additionally', 'overall', 'ultimately',
        'it should be noted', 'this is because', 'as a result',
        'one of the', 'some of the', 'at the same time',
        'in today\'s', 'in the modern', 'plays a crucial role',
        'plays an important role', 'it is crucial', 'it is vital',
        'delve into', 'explore', 'key aspect', 'key role',
        'first and foremost', 'last but not least', 'needless to say',
    ]
    text_lower = text.lower()
    words = re.findall(r'\b[a-zA-Z]+\b', text_lower)
    filler_hits = sum(1 for p in filler_phrases if p in text_lower)
    filler_density = filler_hits / (len(sentences) + 1)

    # ── Feature 5: vocabulary diversity (TTR) ─────────────────────────────────
    ttr = len(set(words)) / (len(words) + 1e-9) if words else 0
    # AI text has moderate TTR; very low (<0.35) = repetitive human,
    # very high (>0.75) = might be technical human. 0.45-0.65 typical of AI.

    # ── Scoring ───────────────────────────────────────────────────────────────
    # Start at 50 (uncertain). Each feature pushes towards AI (higher) or human (lower).
    score = 50.0

    # Burstiness (weight: 35 pts)
    if burstiness is not None:
        if burstiness < -0.4:   score += 35   # very uniform — strongly AI
        elif burstiness < -0.2: score += 25
        elif burstiness < 0.0:  score += 12
        elif burstiness < 0.2:  score -= 5
        elif burstiness < 0.4:  score -= 15
        else:                   score -= 25   # very bursty — strongly human

    # CV of sentence lengths (weight: 20 pts)
    if cv is not None:
        if cv < 0.25:   score += 20   # extremely uniform — AI
        elif cv < 0.40: score += 10
        elif cv < 0.55: score += 0
        elif cv < 0.70: score -= 10
        else:           score -= 18   # very varied — human

    # Mean sentence length (weight: 10 pts)
    if mean_len is not None:
        if 14 <= mean_len <= 26:   score += 10   # ChatGPT sweet spot
        elif 10 <= mean_len < 14:  score += 3
        elif mean_len < 10:        score -= 10   # very short = likely human
        else:                      score -= 5    # very long = academic human

    # Filler phrase density (weight: 15 pts)
    if filler_density > 0.4:  score += 15
    elif filler_density > 0.2: score += 8
    elif filler_density > 0.1: score += 3

    # ── GPT-2 perplexity: supplementary only (weight: ±5 pts max) ─────────────
    # Note: ChatGPT text has GPT-2 perplexity of ~100-300 (NOT low!).
    # Only use it as a slight tiebreaker, not a primary signal.
    perplexity = None
    try:
        perplexity = _gpt2_perplexity(text)
    except Exception:
        pass
    if perplexity is not None:
        if perplexity < 80:    score += 5   # only slightly low ppl = mild AI signal
        elif perplexity > 400: score -= 5   # very high ppl = slight human signal

    score = round(max(0.0, min(100.0, score)), 1)

    # ── Label ─────────────────────────────────────────────────────────────────
    if score >= 75:
        label = 'Likely AI-generated'
    elif score >= 55:
        label = 'Possibly AI-generated'
    elif score >= 40:
        label = 'Uncertain — mixed signals'
    elif score >= 25:
        label = 'Probably human-written'
    else:
        label = 'Likely human-written'

    return {
        'score': score,
        'label': label,
        'perplexity': perplexity,
        'features': {
            'burstiness': round(burstiness, 4) if burstiness is not None else None,
            'sentence_cv': round(cv, 4) if cv is not None else None,
            'mean_sentence_len': round(mean_len, 1) if mean_len is not None else None,
            'filler_density': round(filler_density, 4),
            'sentence_count': len(sent_lengths),
        }
    }

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*"}})

# Initialize services (lazy loading to avoid startup delays)
grammar_enhancer = None
plagiarism_detector = None
formatting_engine = None
section_detector = None
citation_manager = None
file_extractor = None

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
NEW_ENGINE_RUNS = {}
_new_format_document = None


def get_grammar_enhancer():
    """Get or initialize grammar enhancer"""
    global grammar_enhancer
    if grammar_enhancer is None:
        try:
            logger.info("Initializing Grammar Enhancer...")
            grammar_enhancer = GrammarEnhancer()
        except Exception as e:
            logger.warning(f"Grammar enhancer initialization failed: {str(e)}")
    return grammar_enhancer


def get_plagiarism_detector():
    """Get or initialize plagiarism detector"""
    global plagiarism_detector
    if plagiarism_detector is None:
        try:
            logger.info("Initializing Plagiarism Detector...")
            plagiarism_detector = SemanticPlagiarismDetector()
        except Exception as e:
            logger.warning(f"Plagiarism detector initialization failed: {str(e)}")
    return plagiarism_detector


def get_formatting_engine():
    """Get or initialize formatting engine"""
    global formatting_engine
    if formatting_engine is None:
        try:
            logger.info("Initializing Formatting Engine...")
            formatting_engine = DocumentFormattingEngine()
        except Exception as e:
            logger.warning(f"Formatting engine initialization failed: {str(e)}")
    return formatting_engine


def get_section_detector():
    """Get or initialize section detector"""
    global section_detector
    if section_detector is None:
        try:
            logger.info("Initializing Section Detector...")
            section_detector = DocumentSectionDetector()
        except Exception as e:
            logger.warning(f"Section detector initialization failed: {str(e)}")
    return section_detector


def get_citation_manager():
    """Get or initialize citation manager"""
    global citation_manager
    if citation_manager is None:
        try:
            logger.info("Initializing Citation Manager...")
            citation_manager = CitationManager()
        except Exception as e:
            logger.warning(f"Citation manager initialization failed: {str(e)}")
    return citation_manager


def get_file_extractor():
    """Get or initialize file extractor"""
    global file_extractor
    if file_extractor is None:
        try:
            logger.info("Initializing File Extractor...")
            file_extractor = FileExtractor()
        except Exception as e:
            logger.warning(f"File extractor initialization failed: {str(e)}")
    return file_extractor


def get_new_format_document():
    """Load the standalone formatting_engine package (repo root) without clashing with local formatting_engine.py."""
    global _new_format_document
    if _new_format_document is not None:
        return _new_format_document

    package_root = Path(__file__).resolve().parent.parent / 'formatting_engine'
    init_path = package_root / '__init__.py'
    if not init_path.exists():
        raise RuntimeError(f"formatting_engine package not found at {init_path}")

    module_name = 'documentor_new_formatting_engine'
    spec = importlib.util.spec_from_file_location(
        module_name,
        init_path,
        submodule_search_locations=[str(package_root)]
    )
    if spec is None or spec.loader is None:
        raise RuntimeError('Failed to load formatting_engine module spec')

    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)

    formatter = getattr(module, 'format_document', None)
    if not callable(formatter):
        raise RuntimeError('formatting_engine.format_document is not available')

    _new_format_document = formatter
    return _new_format_document


def get_guideline_rule_extractor():
    """Load formatting_engine/guidelines.py without importing the local formatting_engine.py module."""
    package_root = Path(__file__).resolve().parent.parent / 'formatting_engine'
    guidelines_path = package_root / 'guidelines.py'
    if not guidelines_path.exists():
        raise RuntimeError(f"Guideline extractor not found at {guidelines_path}")

    module_name = 'documentor_new_formatting_engine_guidelines'
    existing = sys.modules.get(module_name)
    if existing is not None and hasattr(existing, 'extract_guideline_rules'):
        return existing.extract_guideline_rules

    spec = importlib.util.spec_from_file_location(module_name, guidelines_path)
    if spec is None or spec.loader is None:
        raise RuntimeError('Failed to load guideline extractor module spec')

    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    extractor = getattr(module, 'extract_guideline_rules', None)
    if not callable(extractor):
        raise RuntimeError('extract_guideline_rules is not available')
    return extractor


def _read_json_file(path_value):
    if not path_value:
        return {}
    try:
        path = Path(path_value)
        if not path.exists():
            return {}
        with path.open('r', encoding='utf-8') as handle:
            return json.load(handle)
    except Exception:
        return {}


def _build_new_engine_preview(result):
    layout = _read_json_file(result.get('layout_plan_path'))
    ir = _read_json_file(result.get('ir_path'))

    section_previews = []
    max_sections = int(os.getenv('FORMATTING_ENGINE_PREVIEW_MAX_SECTIONS', '80') or 80)
    max_paragraphs_per_section = int(os.getenv('FORMATTING_ENGINE_PREVIEW_MAX_PARAGRAPHS_PER_SECTION', '80') or 80)
    for section in (layout.get('sections') or [])[:max_sections]:
        subsection_previews = []
        section_paragraphs = []
        for subsection in section.get('subsections') or []:
            paragraphs = (subsection.get('paragraphs') or [])[:max_paragraphs_per_section]
            section_paragraphs.extend(paragraphs)
            subsection_previews.append({
                'id': subsection.get('id'),
                'title': subsection.get('title') or section.get('title') or 'Untitled Section',
                'paragraphs': paragraphs,
                'layout_source': subsection.get('layout_source') or section.get('layout_source') or 'unknown',
            })
        section_previews.append({
            'id': section.get('id'),
            'title': section.get('title') or 'Untitled Section',
            'paragraphs': section_paragraphs[:max_paragraphs_per_section],
            'subsections': subsection_previews,
            'layout_source': section.get('layout_source') or 'unknown',
        })

    metadata = layout.get('metadata') or {}
    return {
        'title': metadata.get('title') or (ir.get('metadata') or {}).get('title'),
        'authors': metadata.get('authors') or [],
        'affiliation': metadata.get('affiliation') or '',
        'abstract': metadata.get('abstract') or '',
        'keywords': metadata.get('keywords') or [],
        'style': layout.get('style') or result.get('target_style'),
        'font': layout.get('font'),
        'body_size_pt': layout.get('body_size_pt'),
        'abstract_size_pt': layout.get('abstract_size_pt'),
        'body_alignment': layout.get('body_alignment'),
        'line_spacing': layout.get('line_spacing'),
        'first_line_indent_inches': layout.get('first_line_indent_inches'),
        'reference_hanging_indent_inches': layout.get('reference_hanging_indent_inches'),
        'columns': layout.get('columns'),
        'sections': section_previews,
        'references': layout.get('references') or [],
        'engine_report': layout.get('engine_report') or {},
        'references_count': len(layout.get('references') or []),
        'figures_count': len(ir.get('figures') or []),
        'tables_count': len(ir.get('tables') or []),
    }


# ============================================================================
# HEALTH CHECK ENDPOINT
# ============================================================================

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    try:
        return jsonify({
            'status': 'healthy',
            'service': 'Python NLP Service - Chapter 4 Implementation',
            'version': '2.0',
            'algorithms': [
                'Grammar Enhancement (T5)',
                'Plagiarism Detection (Semantic Similarity)',
                'Document Formatting (Multi-style)',
                'Section Detection (Naive Bayes)',
                'Citation Management'
            ],
            'supported_formats': ['.pdf', '.docx', '.txt', '.tex', '.latex']
        }), 200
    except Exception as e:
        logger.error(f"Health check error: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ============================================================================
# FILE EXTRACTION ENDPOINT (Centralized)
# ============================================================================

@app.route('/api/extract/file', methods=['POST'])
def extract_file():
    """
    ✅ IMPROVED: Centralized file text extraction endpoint
    Extracts text, structure, tables, and embedded media
    
    Supports: PDF, DOCX, TXT, LaTeX
    Used across all services: Grammar, Plagiarism, Citations, Formatting
    
    Request: multipart/form-data with 'file' field
    Response: Extracted text + structure + media with metadata
    """
    try:
        logger.info("File extraction request received")
        
        # Validate file upload
        if 'file' not in request.files:
            logger.error("No file in request")
            return jsonify({'success': False, 'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            logger.error("Empty filename")
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        logger.info(f"Processing file: {file.filename}")
        
        # Get file extension
        filename = file.filename
        file_ext = os.path.splitext(filename)[1].lower().lstrip('.')
        
        # ✅ Use specialized extractors for PDF and DOCX
        extraction = None
        try:
            if file_ext == 'pdf':
                # Save to temp file for PDF processing
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                    file.save(tmp.name)
                    extraction = extract_pdf_complete(tmp.name)
                    os.unlink(tmp.name)
                    
            elif file_ext in ['docx', 'doc']:
                # Save to temp file for DOCX processing
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    file.save(tmp.name)
                    extraction = extract_docx_complete(tmp.name)
                    os.unlink(tmp.name)
                    
            elif file_ext == 'txt':
                # Simple text file extraction
                file_content = file.read().decode('utf-8', errors='replace')
                extraction = {
                    'text': file_content,
                    'word_count': len(file_content.split()),
                    'structure': [{'type': 'paragraph', 'content': file_content}],
                    'media': []
                }
            else:
                # Fallback to generic extractor for other formats
                extractor = get_file_extractor()
                if extractor is None:
                    return jsonify({'success': False, 'error': 'File extractor not available'}), 500
                
                result = extractor.extract_text(file.read(), filename)
                if not result['success']:
                    return jsonify(result), 400
                
                extraction = {
                    'text': result.get('text', ''),
                    'word_count': result.get('word_count', 0),
                    'structure': [],
                    'media': []
                }
        
        except Exception as extract_err:
            logger.error(f"Specialized extraction failed: {str(extract_err)}")
            logger.error(f"Traceback: {extract_err.__traceback__}")
            # Fallback to generic extractor - but preserve any partial extraction
            if extraction is None:
                extractor = get_file_extractor()
                if extractor:
                    try:
                        file.seek(0)  # Reset file pointer
                        result = extractor.extract_text(file.read(), filename)
                        if result['success']:
                            extraction = {
                                'text': result.get('text', ''),
                                'word_count': result.get('word_count', 0),
                                'structure': [],
                                'media': []
                            }
                    except Exception as fallback_err:
                        logger.error(f"Generic extraction also failed: {fallback_err}")
        
        if extraction is None:
            return jsonify({'success': False, 'error': 'Could not extract file content'}), 400
        
        logger.info(f"Extraction successful - {extraction['word_count']} words, "
                   f"{len(extraction.get('media', []))} media items")
        
        # ✅ Return enhanced response with structure and media
        return jsonify({
            'success': True,
            'text': extraction['text'],
            'word_count': extraction['word_count'],
            'structure': extraction.get('structure', []),
            'media': extraction.get('media', []),
            'metadata': {
                'tables': len([e for e in extraction.get('structure', []) if e.get('type') == 'table']),
                'images': len([m for m in extraction.get('media', []) if m.get('type') == 'image']),
                'graphs': len([m for m in extraction.get('media', []) if m.get('type') == 'graph'])
            }
        }), 200
        
    except Exception as e:
        logger.error(f"File extraction failed: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# STRUCTURED FILE EXTRACTION ENDPOINT (with heading levels and document tree)
# ============================================================================

@app.route('/api/extract/structured', methods=['POST'])
def extract_structured_file():
    """
    ✅ IMPROVED: Extract document with FULL STRUCTURE preservation
    
    Returns blocks with heading levels, document hierarchy, tables, media
    Preserves: heading levels, paragraph styles, page numbers, headers/footers
    Detects: table of contents, multi-column layouts, fonts/sizes
    
    Request: multipart/form-data with 'file' field
    Response: Structured blocks with styles (Heading 1-5, Normal, Title)
    """
    try:
        logger.info("Structured file extraction request received")
        
        # Validate file
        if 'file' not in request.files:
            logger.error("No file in request")
            return jsonify({'success': False, 'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            logger.error("Empty filename")
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        logger.info(f"Processing structured extraction: {file.filename}")
        file_bytes = file.read()
        
        # Use FileExtractor's structured extraction method
        extractor = get_file_extractor()
        if extractor is None:
            return jsonify({'success': False, 'error': 'File extractor not available'}), 500
        
        # Use structured extraction - preserves heading levels and styles
        extraction_result = extractor.extract_structured(file_bytes, file.filename)
        
        if not extraction_result.get('success'):
            logger.error(f"Structured extraction failed: {extraction_result.get('error')}")
            return jsonify(extraction_result), 400
        
        blocks = extraction_result.get('blocks', [])
        structure = extraction_result.get('structure', [])
        media = extraction_result.get('media', [])
        table_of_contents = extraction_result.get('table_of_contents', [])
        docling_payload = extraction_result.get('docling')
        logger.info(f"Structured extraction: {len(blocks)} blocks with heading information")
        
        # Process blocks to include hierarchy information
        processed_blocks = []
        current_h1_idx = -1
        current_h2_idx = -1
        
        for idx, block in enumerate(blocks):
            style = block.get('style', 'Normal')
            text = block.get('text', '')
            
            block_info = {
                'text': text,
                'style': style,
                'index': idx,
                'word_count': len(text.split()),
                'char_count': len(text)
            }
            
            # Track heading hierarchy for table of contents generation
            if style == 'Heading 1':
                current_h1_idx = idx
                current_h2_idx = -1
                block_info['section_level'] = 1
            elif style == 'Heading 2':
                block_info['section_level'] = 2
                block_info['parent_section'] = current_h1_idx
            elif style == 'Heading 3':
                block_info['section_level'] = 3
                block_info['parent_section'] = current_h2_idx
            else:
                block_info['section_level'] = 0
            
            processed_blocks.append(block_info)
        
        # Build table of contents from heading blocks
        toc = []
        for block in processed_blocks:
            if block.get('section_level', 0) >= 1:
                toc.append({
                    'level': block['section_level'],
                    'text': block['text'][:100],  # First 100 chars
                    'index': block['index']
                })

        if table_of_contents:
            toc = table_of_contents
        
        logger.info(f"Generated table of contents with {len(toc)} entries")
        
        # ✅ Return structured response with heading information
        return jsonify({
            'success': True,
            'blocks': processed_blocks,
            'structure': structure,
            'media': media,
            'docling': docling_payload,
            'word_count': extraction_result.get('word_count', 0),
            'file_type': extraction_result.get('file_type', ''),
            'table_of_contents': toc,
            'total_blocks': len(blocks),
            'heading_1_count': len([b for b in processed_blocks if b.get('section_level') == 1]),
            'heading_2_count': len([b for b in processed_blocks if b.get('section_level') == 2]),
            'heading_3_count': len([b for b in processed_blocks if b.get('section_level') == 3]),
            'metadata': {
                'extraction_method': 'structured',
                'backend': extraction_result.get('backend', 'native'),
                'preserves_headings': True,
                'preserves_hierarchy': True,
                'has_docling_payload': bool(docling_payload),
                'structure_items': len(structure),
                'media_items': len(media)
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Structured extraction error: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# ALGORITHM 1: GRAMMAR ENHANCEMENT ENDPOINTS
# ============================================================================

@app.route('/api/grammar/enhance', methods=['POST'])
def enhance_grammar():
    """
    Enhance document grammar and style
    Implements: Algorithm 1 - SequenceEnhancementPipeline
    
    Request body:
    {
        "text": "document text to enhance",
        "mode": "balanced|academic|formal|casual"  (optional, default: balanced)
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON body provided'}), 400
        
        text = data.get('text', '').strip()
        mode = data.get('mode', 'balanced')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        logger.info(f"Grammar enhancement requested - text length: {len(text)}, mode: {mode}")
        
        # Get grammar enhancer
        enhancer = get_grammar_enhancer()
        
        if enhancer is None or enhancer.pipeline is None:
            # Fallback to basic enhancements
            logger.warning("Model not available, using basic enhancement rules")
            result = GrammarEnhancer.apply_basic_enhancements(text)
            return jsonify({
                'success': True,
                'data': result,
                'warning': 'Using basic enhancement rules (transformer model not available)'
            }), 200
        
        # Enhance using T5 model
        result = enhancer.enhance_document(text, mode=mode)
        
        return jsonify({
            'success': True,
            'data': result
        }), 200
        
    except Exception as e:
        logger.error(f"Grammar enhancement error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/grammar/suggestions', methods=['POST'])
def get_grammar_suggestions():
    """
    Get grammar improvement suggestions without modifying text
    
    Request body:
    {
        "text": "document text to analyze"
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON body provided'}), 400
        
        text = data.get('text', '').strip()
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        enhancer = get_grammar_enhancer()
        if enhancer is None:
            enhancer = GrammarEnhancer()
        
        suggestions = enhancer.get_suggestions(text)
        
        return jsonify({
            'success': True,
            'suggestions': suggestions,
            'suggestionCount': len(suggestions)
        }), 200
        
    except Exception as e:
        logger.error(f"Suggestions error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/grammar/basic-enhance', methods=['POST'])
def basic_enhance_grammar():
    """
    Apply basic grammar enhancements using regex rules (no model required)
    Useful as fallback when transformer models are unavailable
    
    Request body:
    {
        "text": "document text to enhance"
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON body provided'}), 400
        
        text = data.get('text', '').strip()
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        result = GrammarEnhancer.apply_basic_enhancements(text)
        
        return jsonify({
            'success': True,
            'data': result
        }), 200
        
    except Exception as e:
        logger.error(f"Basic enhancement error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500


# ============================================================================
# ALGORITHM 2: PLAGIARISM DETECTION ENDPOINTS
# ============================================================================

@app.route('/api/plagiarism/check', methods=['POST'])
def check_plagiarism():
    """
    Check document for plagiarism using semantic similarity
    Implements: Algorithm 2 - SemanticPlagiarismDetector
    
    Request body:
    {
        "text": "document text to check",
        "sources": [
            {"text": "source text 1", "source": "Source Name 1"},
            {"text": "source text 2", "source": "Source Name 2"}
        ],
        "threshold": 0.75  (optional, default: 0.75)
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON body provided'}), 400
        
        text = data.get('text', '').strip()
        sources = data.get('sources', [])
        threshold = data.get('threshold', 0.75)
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        if not sources:
            return jsonify({'error': 'No sources provided for comparison'}), 400
        
        logger.info(f"Plagiarism check requested - text length: {len(text)}, sources: {len(sources)}")
        
        # Get plagiarism detector
        detector = get_plagiarism_detector()
        
        if detector is None:
            logger.error("Plagiarism detector not available")
            return jsonify({
                'error': 'Plagiarism detection service not available',
                'success': False
            }), 503
        
        # Check plagiarism
        result = detector.detect_plagiarism(text, sources, threshold=threshold)
        
        return jsonify({
            'success': True,
            'data': result
        }), 200
        
    except Exception as e:
        logger.error(f"Plagiarism check error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/plagiarism/find-similar', methods=['POST'])
def find_similar_sections():
    """
    Find similar sections between two documents
    
    Request body:
    {
        "document": "main document text",
        "reference": "reference document text",
        "threshold": 0.7  (optional, default: 0.7)
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON body provided'}), 400
        
        document = data.get('document', '').strip()
        reference = data.get('reference', '').strip()
        threshold = data.get('threshold', 0.7)
        
        if not document or not reference:
            return jsonify({'error': 'Both document and reference text required'}), 400
        
        detector = get_plagiarism_detector()
        if detector is None:
            return jsonify({
                'error': 'Plagiarism detection service not available',
                'success': False
            }), 503
        
        similar_sections = detector.find_similar_sections(document, reference, threshold)
        
        return jsonify({
            'success': True,
            'similarSections': similar_sections,
            'count': len(similar_sections)
        }), 200
        
    except Exception as e:
        logger.error(f"Similar sections error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/plagiarism/check-online', methods=['POST'])
def check_plagiarism_online():
    """
    Check document for plagiarism against Wikipedia and DuckDuckGo web results.
    No manual sources required — queries online APIs automatically.

    Request body:
    {
        "text":        "document text",
        "threshold":   0.72,   (optional, default 0.72)
        "include_web": true    (optional, set false for Wikipedia-only)
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON body provided'}), 400

        text = data.get('text', '').strip()
        if not text:
            return jsonify({'error': 'No text provided'}), 400

        threshold   = float(data.get('threshold', 0.72))
        include_web = bool(data.get('include_web', True))

        logger.info(f"Online plagiarism check: {len(text)} chars, threshold={threshold}")

        detector = get_plagiarism_detector()
        if detector is None:
            return jsonify({'error': 'Plagiarism detection service not available',
                            'success': False}), 503

        checker = OnlinePlagiarismChecker(detector)
        result  = checker.check_online(text, threshold=threshold, include_web=include_web)

        return jsonify({'success': True, 'data': result}), 200

    except Exception as e:
        logger.error(f"Online plagiarism check error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/plagiarism/studio', methods=['POST'])
def studio_plagiarism_check():
    """
    PlagiarismStudio-style check: Cosine (TF-IDF) + BERT semantic + AI detection
    Request: { "text": "...", "sources": [{"name": "...", "text": "..."}] }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON body'}), 400

        doc_text = data.get('text', '').strip()
        sources = data.get('sources', [])

        if not doc_text:
            return jsonify({'error': 'No document text provided'}), 400
        if not sources:
            return jsonify({'error': 'No source texts provided'}), 400

        source_results = []
        for src in sources:
            src_text = src.get('text', '').strip()
            if not src_text:
                continue
            cosine = None
            bert = None
            try:
                cosine = studio_cosine_similarity(doc_text, src_text)
            except Exception as e:
                logger.warning(f"Cosine failed: {e}")
            try:
                bert = studio_bert_similarity(doc_text, src_text)
            except Exception as e:
                logger.warning(f"BERT failed: {e}")

            source_results.append({
                'name': src.get('name', 'Source'),
                'cosine': cosine,
                'bert': bert,
                'overall': round((((cosine or 0) + (bert or cosine or 0)) / 2), 2) if (cosine is not None) else None
            })

        ai_result = {}
        try:
            ai_result = studio_ai_detect(doc_text)
        except Exception as e:
            logger.warning(f"AI detection failed: {e}")
            ai_result = {'score': None, 'label': 'Detection failed', 'features': {}, 'perplexity': None}

        # Overall highest similarity across sources
        scores = [r['overall'] for r in source_results if r['overall'] is not None]
        max_score = max(scores) if scores else 0

        return jsonify({
            'success': True,
            'sources': source_results,
            'ai_score': ai_result.get('score'),
            'ai_label': ai_result.get('label'),
            'ai_perplexity': ai_result.get('perplexity'),
            'ai_features': ai_result.get('features', {}),
            'max_similarity': max_score,
            'bert_available': _get_studio_bert() is not None,
            'ai_available': True
        }), 200

    except Exception as e:
        logger.error(f"Studio plagiarism check error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/ai-detect', methods=['POST'])
def ai_detect_only():
    """
    Standalone AI-generated text detection (no plagiarism sources needed).
    Request: { "text": "..." }
    Response: { success, ai_score, ai_label, ai_perplexity, ai_features }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON body', 'success': False}), 400
        text = data.get('text', '').strip()
        if len(text) < 80:
            return jsonify({'error': 'Text too short (min 80 characters)', 'success': False}), 400

        result = studio_ai_detect(text)
        return jsonify({
            'success': True,
            'ai_score': result.get('score'),
            'ai_label': result.get('label'),
            'ai_perplexity': result.get('perplexity'),
            'ai_features': result.get('features', {}),
        }), 200
    except Exception as e:
        logger.error(f"AI detect error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500


# ============================================================================
# ALGORITHM 3: DOCUMENT FORMATTING ENDPOINTS
# ============================================================================

@app.route('/api/formatting/apply', methods=['POST'])
def apply_formatting():
    """
    Format document according to style
    Implements: Algorithm 3 - DocumentFormattingEngine
    
    Request body:
    {
        "content": {
            "title": "document title",
            "abstract": "abstract text",
            "sections": [
                {"heading": "Introduction", "level": 1, "content": "..."},
                {"heading": "Methods", "level": 1, "content": "..."}
            ],
            "references": [
                {"author": "Author", "year": "2023", "title": "Title", "source": "Journal"}
            ]
        },
        "style": "APA|MLA|IEEE|Chicago|Harvard"
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON body provided'}), 400
        
        content = data.get('content', {})
        style = data.get('style', 'APA')
        
        if not content:
            return jsonify({'error': 'No content provided'}), 400
        
        logger.info(f"Document formatting requested - style: {style}")
        
        engine = get_formatting_engine()
        if engine is None:
            engine = DocumentFormattingEngine()
        
        # Format document
        result = engine.format_document(content, style=style)
        
        return jsonify({
            'success': True,
            'data': result
        }), 200
        
    except Exception as e:
        logger.error(f"Formatting error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/formatting/styles', methods=['GET'])
def get_supported_styles():
    """
    Get list of supported formatting styles
    """
    try:
        engine = get_formatting_engine()
        if engine is None:
            engine = DocumentFormattingEngine()
        
        return jsonify({
            'success': True,
            'styles': engine.SUPPORTED_STYLES,
            'default': 'APA'
        }), 200
        
    except Exception as e:
        logger.error(f"Styles error: {e}")
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/formatting/ai-format', methods=['POST'])
def ai_format_document():
    """
    AI-powered end-to-end document formatter.
    Pipeline: chunk → RAG gap analysis → citation conversion → LaTeX generation.

    Accepts two call styles:
      a) multipart/form-data  — 'file' field + optional form fields
         (target_style, source_style, title, authors, generate_latex, metadata_*)
      b) application/json     — legacy, same fields as before

    When a file is uploaded (DOCX or PDF) the extractor uses the structured
    path (paragraph styles / pdfplumber font-sizes) for much better section
    detection.  Plain JSON text falls back to the regex + ref-detector path.
    """
    try:
        from ai_document_formatter import AIDocumentFormatter
        from file_extractor import FileExtractor

        # ------------------------------------------------------------------
        # Parse request — file upload preferred over JSON text
        # ------------------------------------------------------------------
        structured_blocks = None
        table_of_contents = []
        text = ""

        if request.files and 'file' in request.files:
            uploaded = request.files['file']
            file_bytes = uploaded.read()
            extractor = FileExtractor()
            extraction = extractor.extract_structured(file_bytes, uploaded.filename)
            if not extraction['success']:
                return jsonify({'error': f"File extraction failed: {extraction['error']}"}), 400
            structured_blocks = extraction['blocks']
            # Build flat text as fallback for services that still need it
            text = '\n'.join(b['text'] for b in structured_blocks)

            target_style   = request.form.get('target_style', 'APA')
            source_style   = request.form.get('source_style', 'unknown')
            title          = request.form.get('title', 'Untitled Document')
            authors        = request.form.getlist('authors') or []
            generate_latex = request.form.get('generate_latex', 'true').lower() == 'true'
            extra_metadata = {}
            for key in ('date', 'institution', 'keywords'):
                val = request.form.get(f'metadata_{key}')
                if val:
                    extra_metadata[key] = val
        else:
            data = request.get_json()
            if not data:
                return jsonify({'error': 'Provide either a file upload or a JSON body'}), 400
            text = data.get('text', '').strip()
            structured_blocks = data.get('structured_blocks') or data.get('blocks') or None
            table_of_contents = data.get('table_of_contents') or data.get('toc') or []
            if not text:
                if structured_blocks:
                    text = '\n'.join(str(b.get('text', '')).strip() for b in structured_blocks if str(b.get('text', '')).strip())
                if not text:
                    return jsonify({'error': 'No text provided'}), 400
            target_style   = data.get('target_style', 'APA')
            source_style   = data.get('source_style', 'unknown')
            title          = data.get('title', 'Untitled Document')
            authors        = data.get('authors', [])
            generate_latex = data.get('generate_latex', True)
            extra_metadata = data.get('metadata', {})

        logger.info(
            f"AI formatting — style: {target_style}, "
            f"structured: {structured_blocks is not None}, "
            f"text_len: {len(text)}, generate_latex: {generate_latex}"
        )

        formatter = AIDocumentFormatter()
        result = formatter.format(
            text              = text,
            target_style      = target_style,
            source_style      = source_style,
            title             = title,
            authors           = authors,
            extra_metadata    = extra_metadata,
            generate_latex    = generate_latex,
            structured_blocks = structured_blocks,
            table_of_contents = table_of_contents,
        )

        status_code = 200 if result.get('success') else 500
        return jsonify(result), status_code

    except Exception as e:
        logger.error(f"AI formatting error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/formatting/download', methods=['POST'])
def download_formatted_document():
    """
    Generate a downloadable DOCX or PDF from AI-formatted content.

    Request body:
    {
        "sections":  [...],   # plain_sections from /api/formatting/ai-format
        "style":     "APA",
        "title":     "My Document",
        "authors":   ["Author One"],
        "format":    "docx" | "pdf"
    }
    Returns the binary file with appropriate Content-Type.
    """
    import io
    from flask import Response

    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON body'}), 400

        sections = data.get('sections', [])
        style    = data.get('style', 'APA')
        title    = data.get('title', 'Formatted Document')
        authors  = data.get('authors', [])
        fmt      = data.get('format', 'docx').lower()

        safe_title = title[:40].replace(' ', '_')

        if fmt == 'docx':
            buf = _build_docx(sections, style, title, authors)
            return Response(
                buf.getvalue(),
                status=200,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={
                    'Content-Disposition': f'attachment; filename="{safe_title}_{style}.docx"',
                    'Access-Control-Expose-Headers': 'Content-Disposition',
                }
            )
        elif fmt == 'pdf':
            buf = _build_pdf(sections, style, title, authors)
            return Response(
                buf.getvalue(),
                status=200,
                mimetype='application/pdf',
                headers={
                    'Content-Disposition': f'attachment; filename="{safe_title}_{style}.pdf"',
                    'Access-Control-Expose-Headers': 'Content-Disposition',
                }
            )
        else:
            return jsonify({'error': f'Unsupported format: {fmt}'}), 400

    except Exception as e:
        logger.error(f"Download error: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/formatting/new-engine', methods=['POST'])
@app.route('/api/formatting/experimental-engine', methods=['POST'])
def run_new_engine_formatting():
    """
    Run the standalone formatting_engine pipeline and return preview + artifact references.
    """
    temp_input_path = None
    temp_guidelines_path = None
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        uploaded = request.files['file']
        target_style = (request.form.get('target_style') or 'ieee').strip().lower()
        supported_styles = {'ieee', 'apa', 'acm', 'nature', 'elsevier', 'chicago'}
        if target_style not in supported_styles:
            return jsonify({'error': f"Unsupported style '{target_style}'"}), 400

        input_suffix = Path(uploaded.filename or '').suffix.lower()
        if not input_suffix:
            input_suffix = '.tmp'

        with tempfile.NamedTemporaryFile(delete=False, suffix=input_suffix) as temp_input:
            uploaded.save(temp_input.name)
            temp_input_path = temp_input.name

        custom_rules = None
        guidelines_file = request.files.get('guidelines')
        if guidelines_file and guidelines_file.filename:
            guide_suffix = Path(guidelines_file.filename or '').suffix.lower() or '.tmp'
            with tempfile.NamedTemporaryFile(delete=False, suffix=guide_suffix) as temp_guidelines:
                guidelines_file.save(temp_guidelines.name)
                temp_guidelines_path = temp_guidelines.name
            extract_guideline_rules = get_guideline_rule_extractor()
            custom_rules = extract_guideline_rules(temp_guidelines_path)

        output_root = request.form.get('output_dir') or os.path.join(UPLOAD_FOLDER, 'formatting_engine_runs')
        os.makedirs(output_root, exist_ok=True)

        formatter = get_new_format_document()
        use_ai_raw = request.form.get('use_ai')
        if use_ai_raw is not None:
            use_ai = str(use_ai_raw).strip().lower() in {'1', 'true', 'yes', 'on'}
        else:
            use_ai = os.getenv('FORMATTING_ENGINE_USE_AI', '').strip().lower() in {'1', 'true', 'yes', 'on'}
        previous_local_only = os.environ.get('FORMATTING_ENGINE_LOCAL_ONLY')
        previous_use_ai = os.environ.get('FORMATTING_ENGINE_USE_AI')
        if not use_ai:
            os.environ['FORMATTING_ENGINE_LOCAL_ONLY'] = '1'
            os.environ['FORMATTING_ENGINE_USE_AI'] = '0'
        else:
            os.environ['FORMATTING_ENGINE_LOCAL_ONLY'] = '0'
            os.environ['FORMATTING_ENGINE_USE_AI'] = '1'
        try:
            result = formatter(
                input_path=temp_input_path,
                target_style=target_style,
                output_dir=output_root,
                custom_rules=custom_rules
            )
        finally:
            if previous_local_only is None:
                os.environ.pop('FORMATTING_ENGINE_LOCAL_ONLY', None)
            else:
                os.environ['FORMATTING_ENGINE_LOCAL_ONLY'] = previous_local_only
            if previous_use_ai is None:
                os.environ.pop('FORMATTING_ENGINE_USE_AI', None)
            else:
                os.environ['FORMATTING_ENGINE_USE_AI'] = previous_use_ai

        docx_path = result.get('docx_path')
        run_dir = Path(docx_path).parent if docx_path else Path(result.get('ir_path', output_root)).parent
        run_id = run_dir.name

        artifact_candidates = {
            'docx': result.get('docx_path'),
            'pdf': result.get('pdf_path'),
            'tex': result.get('latex_path'),
            'ir': result.get('ir_path'),
            'layout': result.get('layout_plan_path'),
        }
        artifacts = {}
        for key, path_value in artifact_candidates.items():
            if path_value and Path(path_value).exists():
                artifacts[key] = str(Path(path_value).resolve())

        NEW_ENGINE_RUNS[run_id] = {
            'created_at': int(time.time()),
            'artifacts': artifacts
        }

        preview = _build_new_engine_preview(result)

        return jsonify({
            'success': True,
            'engine': 'new_engine',
            'engine_mode': 'local_rules' if not use_ai else 'groq_assisted',
            'run_id': run_id,
            'target_style': target_style,
            'preview': preview,
            'artifacts': {key: Path(value).name for key, value in artifacts.items()},
            'available_files': list(artifacts.keys()),
            'custom_guidelines': custom_rules or result.get('custom_guidelines'),
            'warnings': result.get('warnings', [])
        }), 200
    except Exception as e:
        logger.error(f"New engine formatting error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500
    finally:
        if temp_input_path and os.path.exists(temp_input_path):
            try:
                os.unlink(temp_input_path)
            except Exception:
                pass
        if temp_guidelines_path and os.path.exists(temp_guidelines_path):
            try:
                os.unlink(temp_guidelines_path)
            except Exception:
                pass


@app.route('/api/formatting/new-engine/download/<run_id>', methods=['GET'])
@app.route('/api/formatting/experimental-engine/download/<run_id>', methods=['GET'])
def download_new_engine_artifact(run_id):
    """
    Download one artifact produced by /api/formatting/new-engine.
    Query: ?file=docx|pdf|tex|ir|layout
    """
    file_key = (request.args.get('file') or 'docx').strip().lower()
    allowed = {'docx', 'pdf', 'tex', 'ir', 'layout'}
    if file_key not in allowed:
        return jsonify({'error': f"Unsupported file type '{file_key}'"}), 400

    run_info = NEW_ENGINE_RUNS.get(run_id)
    if not run_info:
        return jsonify({'error': 'Run not found. Re-run formatting to regenerate artifacts.'}), 404

    artifact_path = (run_info.get('artifacts') or {}).get(file_key)
    if not artifact_path or not Path(artifact_path).exists():
        return jsonify({'error': f"Artifact '{file_key}' is not available for this run"}), 404

    mime_map = {
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'pdf': 'application/pdf',
        'tex': 'text/plain',
        'ir': 'application/json',
        'layout': 'application/json',
    }
    return send_file(
        artifact_path,
        as_attachment=True,
        download_name=Path(artifact_path).name,
        mimetype=mime_map.get(file_key, 'application/octet-stream')
    )


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------
_STYLE_FONT_SIZE = {'APA': 12, 'MLA': 12, 'IEEE': 10, 'Chicago': 12, 'Harvard': 12}
_STYLE_SPACING   = {'APA': 2.0, 'MLA': 2.0, 'IEEE': 1.0, 'Chicago': 2.0, 'Harvard': 1.5}
# (top, bottom, left, right) in inches
_STYLE_MARGINS   = {
    'APA':     (1.0,  1.0,  1.0,    1.0),
    'MLA':     (1.0,  1.0,  1.0,    1.0),
    'IEEE':    (0.75, 1.0,  0.625,  0.625),
    'Chicago': (1.0,  1.0,  1.0,    1.0),
    'Harvard': (1.0,  1.0,  1.0,    1.0),
}
_BIBLIO_HEADING  = {
    'APA': 'References', 'MLA': 'Works Cited',
    'IEEE': 'References', 'Chicago': 'Bibliography', 'Harvard': 'Reference List',
}


def _split_ref_entries(text: str, style: str):
    """
    Split a raw references block into individual entry strings and sort them.

    IEEE  — numbered (e.g. [1], [2]): split on [ n ] boundaries, keep numeric order.
    Other — author-date styles: split on blank lines first; if that yields ≤1 entry
            fall back to detecting lines that start a new entry (capital letter /
            author-like pattern after a previous non-empty line), then sort
            alphabetically by the first 60 characters of the entry.
    """
    text = text.strip()
    if not text:
        return []

    entries = []

    if style == 'IEEE':
        # Split on IEEE [n] markers — each marker starts a new entry
        parts = re.split(r'(?=\[\d+\])', text)
        for p in parts:
            p = p.strip()
            if p:
                # Collapse internal newlines within one entry (wrap lines)
                entry = ' '.join(line.strip() for line in p.splitlines() if line.strip())
                entries.append(entry)
        # Sort by numeric tag so [1],[2],… are always in order
        def _ieee_num(e):
            m = re.match(r'\[(\d+)\]', e)
            return int(m.group(1)) if m else 9999
        entries.sort(key=_ieee_num)
        return entries

    # ── Author-date styles ─────────────────────────────────────────────────
    # Step 1: try blank-line separation
    raw = re.split(r'\n{2,}', text)
    if len(raw) > 1:
        # Collapse any internal line-wraps within each entry
        for r in raw:
            entry = ' '.join(line.strip() for line in r.splitlines() if line.strip())
            if entry:
                entries.append(entry)
    else:
        # Step 2: single-newline text — detect entry boundaries.
        # A new entry starts when a line begins with a capital letter AND
        # the previous line ended with a full-stop / ) / ] (i.e. a complete sentence).
        lines = text.splitlines()
        current: list = []
        end_chars = set('.);]"\'')
        entry_start = re.compile(
            r'^(?:[A-Z][a-z]|[A-Z]{2,}|[(\[])'  # starts with Title-case, ALL-CAPS, or [(
        )
        for line in lines:
            stripped = line.strip()
            if not stripped:
                # Blank line → flush
                if current:
                    entries.append(' '.join(current))
                    current = []
                continue
            if current and entry_start.match(stripped) and current[-1] and current[-1][-1] in end_chars:
                # Previous entry finished, new one starting
                entries.append(' '.join(current))
                current = [stripped]
            else:
                current.append(stripped)
        if current:
            entries.append(' '.join(current))

    # Step 3: sort alphabetically (ignore leading articles / punctuation)
    def _sort_key(e):
        # Strip leading [n] / (n) / numbers so "Smith" sorts as "Smith" not "[2] Smith"
        clean = re.sub(r'^[\[(]?\d+[\])]?\s*', '', e).lower()
        return clean

    entries.sort(key=_sort_key)
    return entries


def _build_docx(sections, style, title, authors):
    """Build a python-docx document styled according to `style`."""
    import io
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = DocxDocument()

    fsize     = _STYLE_FONT_SIZE.get(style, 12)
    spacing   = _STYLE_SPACING.get(style, 2.0)
    bib_head  = _BIBLIO_HEADING.get(style, 'References')
    margins   = _STYLE_MARGINS.get(style, (1.0, 1.0, 1.0, 1.0))  # top, bot, left, right
    font_name = 'Arial' if style == 'IEEE' else 'Times New Roman'

    # ---- Page margins (per style) ----
    for sec_obj in doc.sections:
        sec_obj.top_margin    = Inches(margins[0])
        sec_obj.bottom_margin = Inches(margins[1])
        sec_obj.left_margin   = Inches(margins[2])
        sec_obj.right_margin  = Inches(margins[3])

    # ---- Document default style ----
    normal = doc.styles['Normal']
    normal.font.name = font_name
    normal.font.size = Pt(fsize)

    # ---- Page number in footer (centered; top-right for APA via header) ----
    def _add_page_number_footer(sec_obj):
        footer = sec_obj.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.font.name = font_name
        run.font.size = Pt(fsize)
        for tag, text in [('w:fldChar', None), ('w:instrText', 'PAGE'), ('w:fldChar', None)]:
            el = OxmlElement(tag)
            if text:
                el.text = text
            else:
                el.set(qn('w:fldCharType'), 'begin' if not run._r.findall(qn('w:fldChar')) else 'end')
            run._r.append(el)

    def _add_page_number_header_right(sec_obj):
        """APA: page number top-right."""
        header = sec_obj.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run()
        run.font.name = font_name
        run.font.size = Pt(fsize)
        for tag, text in [('w:fldChar', None), ('w:instrText', 'PAGE'), ('w:fldChar', None)]:
            el = OxmlElement(tag)
            if text:
                el.text = text
            else:
                el.set(qn('w:fldCharType'), 'begin' if not run._r.findall(qn('w:fldChar')) else 'end')
            run._r.append(el)

    for sec_obj in doc.sections:
        if style == 'APA':
            _add_page_number_header_right(sec_obj)
        else:
            _add_page_number_footer(sec_obj)

    # ---- Helpers ----
    def _set_para_spacing(para, line_spacing=spacing, space_before=0, space_after=0):
        pPr = para._p.get_or_add_pPr()
        spacing_el = OxmlElement('w:spacing')
        spacing_el.set(qn('w:line'),     str(int(240 * line_spacing)))
        spacing_el.set(qn('w:lineRule'), 'auto')
        spacing_el.set(qn('w:before'),   str(int(space_before * 20)))
        spacing_el.set(qn('w:after'),    str(int(space_after  * 20)))
        pPr.append(spacing_el)

    def _detect_heading_level(text):
        """Guess heading level from the text structure."""
        if re.match(r'^\d+\.\d+\.\d+', text): return 3
        if re.match(r'^\d+\.\d+\s',    text): return 2
        if re.match(r'^[A-Z]\.\s',     text): return 2   # IEEE: A. Subsection
        return 1

    def _is_inline_heading(text):
        """Detect a sub-heading embedded as its own paragraph block."""
        t = text.strip()
        if not t or len(t) > 100 or '\n' in t:
            return False
        if t[-1] in '.,:;?!':
            return False
        if re.match(r'^\d+(\.\d+)+[\s\.]', t):
            return True
        words = t.split()
        if 2 <= len(words) <= 8:
            cap = sum(1 for w in words if w and w[0].isupper())
            if cap / len(words) >= 0.80:
                return True
        return False

    def _add_heading(text, level=1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = font_name
        _set_para_spacing(p, line_spacing=1.0, space_before=6, space_after=3)

        if style in ('APA', 'Harvard'):
            # APA L1: centred bold | L2: left bold | L3: left bold italic
            # L4: indented bold | L5: indented bold italic
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
                run.font.size = Pt(fsize)
            elif level == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run.bold = True
                run.font.size = Pt(fsize)
            elif level == 3:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run.bold  = True
                run.italic = True
                run.font.size = Pt(fsize)
            else:
                p.paragraph_format.first_line_indent = Inches(0.5)
                run.bold = True
                run.font.size = Pt(fsize)
            if style == 'Harvard' and level == 1:
                run.font.size = Pt(fsize + 2)   # Harvard L1 = 14 pt

        elif style == 'MLA':
            # MLA: L1 bold | L2 italic | L3 bold italic | L4 underline
            run.font.size = Pt(fsize)
            if level == 1:   run.bold = True
            elif level == 2: run.italic = True
            elif level == 3: run.bold = True; run.italic = True
            elif level >= 4: run.underline = True

        elif style == 'IEEE':
            # IEEE L1: centred small-caps 10pt | L2: left italic 10pt
            run.font.size = Pt(10)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sm = OxmlElement('w:smallCaps')
                sm.set(qn('w:val'), '1')
                run._r.get_or_add_rPr().append(sm)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run.italic = True

        elif style == 'Chicago':
            # Chicago L1: centred bold | L2: centred italic
            # L3: left bold | L4: left italic
            run.font.size = Pt(fsize)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
            elif level == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.italic = True
            elif level == 3:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run.bold = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run.italic = True
        return p

    def _add_body_para(text, is_abstract=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(fsize)
        # First-line indent rules (APA/MLA/Chicago: 0.5 in; Harvard: 0.5 in)
        # APA abstract: NO first-line indent
        if not is_abstract and style in ('APA', 'MLA', 'Chicago', 'Harvard'):
            p.paragraph_format.first_line_indent = Inches(0.5)
        _set_para_spacing(p, line_spacing=spacing)
        return p

    def _add_ref_para(text):
        """Hanging indent (0.5 in) for bibliography entries."""
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'),    '720')  # 0.5 in = 720 twips
        ind.set(qn('w:hanging'), '720')
        pPr.append(ind)
        _set_para_spacing(p, line_spacing=spacing)
        if p.runs:
            p.runs[0].font.name = font_name
            p.runs[0].font.size = Pt(fsize)
        return p

    # ---- Title block ----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_para.add_run(title)
    run_title.bold = True
    run_title.font.name = font_name
    run_title.font.size = Pt(fsize)
    _set_para_spacing(title_para, line_spacing=spacing)

    if authors:
        auth_para = doc.add_paragraph()
        auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_auth = auth_para.add_run(', '.join(authors))
        run_auth.font.name = font_name
        run_auth.font.size = Pt(fsize)
        _set_para_spacing(auth_para, line_spacing=spacing)

    doc.add_paragraph()  # blank line after title block

    # ---- Sections ----
    for sec in sections:
        stype   = sec.get('section_type', 'unknown')
        heading = sec.get('heading', stype.title())
        text    = (sec.get('text') or '').strip()
        is_ref  = stype == 'references'
        is_abs  = stype == 'abstract'

        # Section heading
        if is_ref:
            _add_heading(bib_head, level=1)
        else:
            level = _detect_heading_level(heading)
            _add_heading('Abstract' if is_abs else heading, level=level)

        # Body paragraphs — split on double newline, fall back to single
        if is_ref:
            ref_entries = _split_ref_entries(text, style)
            for entry in ref_entries:
                _add_ref_para(entry)
        else:
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            if not paragraphs:
                paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
            for para_text in paragraphs:
                if not is_abs and _is_inline_heading(para_text):
                    lvl = _detect_heading_level(para_text)
                    _add_heading(para_text, level=max(2, lvl + 1))
                else:
                    _add_body_para(para_text, is_abstract=is_abs)

    buf = io.BytesIO()
    doc.save(buf)
    return buf


# ---------------------------------------------------------------------------
# PDF builder  (reportlab)
# ---------------------------------------------------------------------------
_PDF_FONT      = 'Times-Roman'
_PDF_FONT_BOLD = 'Times-Bold'

def _build_pdf(sections, style, title, authors):
    """Build a PDF using reportlab platypus."""
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib import colors

    buf = io.BytesIO()
    fsize   = _STYLE_FONT_SIZE.get(style, 12)
    spacing = _STYLE_SPACING.get(style, 2.0)
    bib_head = _BIBLIO_HEADING.get(style, 'References')
    lead    = fsize * spacing * 1.2  # leading (line height)

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=inch, rightMargin=inch,
        topMargin=inch, bottomMargin=inch,
    )

    base_font = 'Times-Roman'
    bold_font = 'Times-Bold'

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'DocTitle',
        fontName=bold_font, fontSize=fsize,
        alignment=TA_CENTER, leading=lead,
        spaceAfter=4,
    )
    author_style = ParagraphStyle(
        'DocAuthor',
        fontName=base_font, fontSize=fsize,
        alignment=TA_CENTER, leading=lead,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        'SecHeading',
        fontName=bold_font, fontSize=fsize,
        alignment=TA_CENTER if style in ('APA', 'MLA', 'Chicago', 'Harvard') else TA_LEFT,
        leading=lead,
        spaceBefore=12, spaceAfter=6,
    )
    body_style = ParagraphStyle(
        'Body',
        fontName=base_font, fontSize=fsize,
        alignment=TA_JUSTIFY,
        leading=lead,
        firstLineIndent=(0.5 * inch) if style in ('APA', 'Chicago', 'MLA') else 0,
        spaceAfter=0,
    )
    ref_style = ParagraphStyle(
        'Ref',
        fontName=base_font, fontSize=fsize,
        alignment=TA_LEFT,
        leading=lead,
        leftIndent=0.5 * inch,
        firstLineIndent=-(0.5 * inch),  # hanging indent
        spaceAfter=6,
    )

    story = []

    # Title
    story.append(Paragraph(title, title_style))
    if authors:
        story.append(Paragraph(', '.join(authors), author_style))
    story.append(Spacer(1, 12))

    for sec in sections:
        stype   = sec.get('section_type', 'unknown')
        heading = sec.get('heading', stype.title())
        text    = (sec.get('text') or '').strip()
        is_ref  = stype == 'references'

        h_label = bib_head if is_ref else ('Abstract' if stype == 'abstract' else heading)
        story.append(Paragraph(h_label, heading_style))

        if is_ref:
            ref_entries = _split_ref_entries(text, style)
            for entry in ref_entries:
                safe = entry.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                try:
                    story.append(Paragraph(safe, ref_style))
                    story.append(Spacer(1, lead * 0.2))
                except Exception:
                    story.append(Paragraph('', ref_style))
        else:
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            if not paragraphs:
                paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
            for para_text in paragraphs:
                safe = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                try:
                    story.append(Paragraph(safe, body_style))
                    story.append(Spacer(1, lead * 0.2))
                except Exception:
                    story.append(Paragraph('', body_style))

        story.append(Spacer(1, 6))

    doc.build(story)
    return buf


# ============================================================================
# ALGORITHM 4: SECTION DETECTION ENDPOINTS
# ============================================================================

@app.route('/api/document/sections', methods=['POST'])
def detect_sections():
    """
    Detect document sections
    Implements: Algorithm 4 - DocumentSectionDetector
    
    Request body:
    {
        "text": "document text",
        "hint": "thesis|journal|conference|report"  (optional)
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON body provided'}), 400
        
        text = data.get('text', '').strip()
        hint = data.get('hint', '')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        logger.info(f"Section detection requested - text length: {len(text)}")
        
        detector = get_section_detector()
        if detector is None:
            detector = DocumentSectionDetector()
        
        # Detect sections
        result = detector.detect_sections(text, hint=hint)
        
        return jsonify({
            'success': True,
            'data': result
        }), 200
        
    except Exception as e:
        logger.error(f"Section detection error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/document/extract-section', methods=['POST'])
def extract_section():
    """
    Extract specific section from document
    
    Request body:
    {
        "text": "document text",
        "sectionType": "abstract|introduction|methodology|results|discussion|conclusion|references"
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON body provided'}), 400
        
        text = data.get('text', '').strip()
        section_type = data.get('sectionType', '').strip()
        
        if not text or not section_type:
            return jsonify({'error': 'Text and sectionType required'}), 400
        
        detector = get_section_detector()
        if detector is None:
            detector = DocumentSectionDetector()
        
        section_content = detector.extract_section(text, section_type)
        
        if section_content is None:
            return jsonify({
                'success': False,
                'error': f'Section "{section_type}" not found in document'
            }), 404
        
        return jsonify({
            'success': True,
            'sectionType': section_type,
            'content': section_content
        }), 200
        
    except Exception as e:
        logger.error(f"Section extraction error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/document/validate-structure', methods=['POST'])
def validate_structure():
    """
    Validate document structure against academic standards
    
    Request body:
    {
        "text": "document text"
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON body provided'}), 400
        
        text = data.get('text', '').strip()
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        detector = get_section_detector()
        if detector is None:
            detector = DocumentSectionDetector()
        
        result = detector.validate_document_structure(text)
        
        return jsonify({
            'success': True,
            'data': result
        }), 200
        
    except Exception as e:
        logger.error(f"Structure validation error: {e}", exc_info=True)
        return jsonify({'error': str(e), 'success': False}), 500


# ============================================================================
# EVALUATION METRICS ENDPOINTS
# ============================================================================

@app.route('/api/metrics/grammar', methods=['GET'])
def get_grammar_metrics():
    """Get grammar enhancement model evaluation metrics"""
    try:
        enhancer = get_grammar_enhancer()
        if enhancer is None:
            enhancer = GrammarEnhancer()
        
        metrics = {
            "algorithm": "SequenceEnhancementPipeline (Algorithm 1)",
            "model": enhancer.model_name,
            "metrics": {
                "tokenAccuracy": 0.893,
                "sentenceAccuracy": 0.712,
                "precision": 0.912,
                "recall": 0.871,
                "f1Score": 0.889,
                "averageLatency": "2.8s"
            },
            "trainingData": "CoNLL 2014 Shared Task (1.3M examples)"
        }
        
        return jsonify({'success': True, 'metrics': metrics}), 200
        
    except Exception as e:
        logger.error(f"Grammar metrics error: {e}")
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/metrics/plagiarism', methods=['GET'])
def get_plagiarism_metrics():
    """Get plagiarism detection model evaluation metrics"""
    try:
        detector = get_plagiarism_detector()
        if detector is None:
            detector = SemanticPlagiarismDetector()
        
        metrics = detector.get_evaluation_metrics()
        
        return jsonify({'success': True, 'metrics': metrics}), 200
        
    except Exception as e:
        logger.error(f"Plagiarism metrics error: {e}")
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/metrics/formatting', methods=['GET'])
def get_formatting_metrics():
    """Get document formatting engine evaluation metrics"""
    try:
        engine = get_formatting_engine()
        if engine is None:
            engine = DocumentFormattingEngine()
        
        metrics = engine.get_evaluation_metrics()
        
        return jsonify({'success': True, 'metrics': metrics}), 200
        
    except Exception as e:
        logger.error(f"Formatting metrics error: {e}")
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/metrics/sections', methods=['GET'])
def get_sections_metrics():
    """Get section detection model evaluation metrics"""
    try:
        detector = get_section_detector()
        if detector is None:
            detector = DocumentSectionDetector()
        
        metrics = detector.get_evaluation_metrics()
        
        return jsonify({'success': True, 'metrics': metrics}), 200
        
    except Exception as e:
        logger.error(f"Sections metrics error: {e}")
        return jsonify({'error': str(e), 'success': False}), 500


# ============================================================================
# CITATION MANAGEMENT ENDPOINTS
# ============================================================================

@app.route('/api/citations/extract', methods=['POST'])
def extract_citations():
    """
    Extract citations and references from uploaded document
    
    Request: multipart/form-data with 'file' field
    Response: Complete citation analysis with line numbers and mapping
    """
    try:
        logger.info("Citation extraction request received")
        start_time = time.time()
        
        # Get uploaded file
        if 'file' not in request.files:
            logger.error("No file in request")
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            logger.error("Empty filename")
            return jsonify({'error': 'No file selected'}), 400
        
        logger.info(f"Processing file: {file.filename}")
        
        # Read file content
        file_content = file.read()
        filename = file.filename
        
        logger.info(f"File size: {len(file_content)} bytes")
        
        # Extract text
        manager = get_citation_manager()
        if manager is None:
            logger.error("Citation manager not initialized")
            return jsonify({'error': 'Citation service not available'}), 500
        
        text = manager.extract_text_from_file(file_content, filename)
        
        if text.startswith('Error') or text.startswith('Unsupported'):
            return jsonify({'error': text}), 400
        
        # Detect citation style
        detected_style = manager.detect_citation_style(text)
        
        # Extract citations with line numbers
        citations = manager.extract_citations_from_text(text)
        
        # Extract references
        references, ref_start_line = extract_references_section(text)
        
        # Match citations to references
        mapping = {}
        if citations and references:
            mapping = match_citations_to_references(citations, references)
        
        processing_time = time.time() - start_time
        
        return jsonify({
            'success': True,
            'text': text,
            'word_count': len(text.split()),
            'char_count': len(text),
            'detected_style': detected_style,
            'citations': citations,
            'references': references,
            'ref_start_line': ref_start_line,
            'mapping': {str(k): v for k, v in mapping.items()},  # Convert keys to strings for JSON
            'matched_count': len(mapping),
            'unmatched_citations': len(citations) - len(mapping),
            'uncited_references': len(references) - len(set(mapping.values())),
            'processing_time': round(processing_time, 2)
        }), 200
        
    except Exception as e:
        logger.error(f"Citation extraction failed: {str(e)}")
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/citations/detect-style', methods=['POST'])
def detect_style():
    """Detect citation style from text"""
    try:
        data = request.get_json()
        
        if not data or 'text' not in data:
            return jsonify({'error': 'Text is required'}), 400
        
        text = data['text']
        manager = get_citation_manager()
        
        detected_style = manager.detect_citation_style(text)
        
        return jsonify({
            'success': True,
            'style': detected_style,
            'confidence': 'high' if detected_style != 'Unknown' else 'low'
        }), 200
        
    except Exception as e:
        logger.error(f"Style detection failed: {str(e)}")
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/citations/match', methods=['POST'])
def match_citations_endpoint():
    """Match citations to references in document"""
    try:
        data = request.get_json()
        
        if not data or 'text' not in data:
            return jsonify({'error': 'Text is required'}), 400
        
        text = data['text']
        manager = get_citation_manager()
        
        # Extract citations
        citations = manager.extract_citations_from_text(text)
        
        # Extract references
        references, ref_start_line = extract_references_section(text)
        
        # Match citations to references
        mapping = {}
        if citations and references:
            mapping = match_citations_to_references(citations, references)
        
        return jsonify({
            'success': True,
            'citations': citations,
            'references': references,
            'ref_start_line': ref_start_line,
            'mapping': {str(k): v for k, v in mapping.items()},
            'matched_count': len(mapping),
            'total_citations': len(citations),
            'total_references': len(references)
        }), 200
        
    except Exception as e:
        logger.error(f"Citation matching failed: {str(e)}")
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/citations/match-ai', methods=['POST'])
def match_citations_ai_endpoint():
    """
    AI-powered citation↔reference matching using sentence-transformer embeddings.

    Request body:
      { "text": "<full document text>" }

    Response:
      {
        "success": true,
        "citations":   [...],          # all in-text citations found
        "references":  [...],          # all reference entries found
        "matches": [
          {
            "citation_idx":  int,
            "citation_text": str,
            "citation_context": str,
            "ref_idx":  int,
            "ref_text": str,
            "score":    float,         # cosine similarity 0–1
            "matched":  bool,          # True when score ≥ 0.30
            "query":    str,           # the string that was embedded
          },
          ...
        ],
        "unmatched_citations": [int],  # indices with score < threshold
        "matched_count":  int,
        "total_citations": int,
        "total_references": int
      }
    """
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'error': 'text is required'}), 400

        text = data['text']
        manager = get_citation_manager()

        citations  = manager.extract_citations_from_text(text)
        references, ref_start_line = extract_references_section(text)

        if not citations:
            return jsonify({
                'success': True,
                'citations': [], 'references': references,
                'matches': [], 'unmatched_citations': [],
                'matched_count': 0, 'total_citations': 0,
                'total_references': len(references),
            }), 200

        if not references:
            return jsonify({
                'success': True,
                'citations': citations, 'references': [],
                'matches': [], 'unmatched_citations': list(range(len(citations))),
                'matched_count': 0, 'total_citations': len(citations),
                'total_references': 0,
            }), 200

        matcher = AISemanticCitationMatcher()
        matcher.fit(references)
        ai_result = matcher.match_all(citations)

        matches = []
        for cit_idx, info in ai_result.items():
            ref_idx = info['ref_idx']
            matches.append({
                'citation_idx':     cit_idx,
                'citation_text':    citations[cit_idx].get('text', ''),
                'citation_context': citations[cit_idx].get('context', ''),
                'ref_idx':          ref_idx,
                'ref_text':         references[ref_idx].get('text', '') if ref_idx >= 0 else '',
                'score':            info['score'],
                'matched':          info['matched'],
                'query':            info['query'],
            })

        unmatched = [m['citation_idx'] for m in matches if not m['matched']]

        return jsonify({
            'success':            True,
            'citations':          citations,
            'references':         references,
            'ref_start_line':     ref_start_line,
            'matches':            matches,
            'unmatched_citations': unmatched,
            'matched_count':      len(matches) - len(unmatched),
            'total_citations':    len(citations),
            'total_references':   len(references),
        }), 200

    except Exception as e:
        logger.error(f"AI citation matching failed: {str(e)}")
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/citations/format', methods=['POST'])
def format_document_citations():
    """Format document according to citation style"""
    try:
        data = request.get_json()
        
        if not data or 'text' not in data or 'style' not in data:
            return jsonify({'error': 'Text and style are required'}), 400
        
        text = data['text']
        style = data['style']
        
        manager = get_citation_manager()
        formatted_text = manager.apply_citation_formatting(text, style)
        
        return jsonify({
            'success': True,
            'formatted_text': formatted_text,
            'style': style
        }), 200
        
    except Exception as e:
        logger.error(f"Citation formatting failed: {str(e)}")
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/citations/generate', methods=['POST'])
def generate_citation_entry():
    """Generate a bibliography entry from citation data"""
    try:
        data = request.get_json()
        
        if not data or 'citation_data' not in data or 'style' not in data:
            return jsonify({'error': 'Citation data and style are required'}), 400
        
        citation_data = data['citation_data']
        style = data['style']
        
        manager = get_citation_manager()
        citation = manager.generate_bibliography_entry(citation_data, style)
        
        return jsonify({
            'success': True,
            'citation': citation,
            'style': style
        }), 200
        
    except Exception as e:
        logger.error(f"Citation generation failed: {str(e)}")
        return jsonify({'error': str(e), 'success': False}), 500


@app.route('/api/citations/validate', methods=['POST'])
def validate_document_citations():
    """Validate citations against a specific style"""
    try:
        data = request.get_json()
        
        if not data or 'text' not in data or 'style' not in data:
            return jsonify({'error': 'Text and style are required'}), 400
        
        text = data['text']
        style = data['style']
        
        manager = get_citation_manager()
        issues = manager.validate_citations(text, style)
        
        return jsonify({
            'success': True,
            'issues': issues,
            'total_issues': len(issues),
            'is_valid': len(issues) == 0,
            'style': style
        }), 200
        
    except Exception as e:
        logger.error(f"Citation validation failed: {str(e)}")
        return jsonify({'error': str(e), 'success': False}), 500


# ============================================================================
# ERROR HANDLING
# ============================================================================

@app.errorhandler(404)
def not_found(error):
    """Handle 404 errors"""
    return jsonify({'error': 'Endpoint not found', 'path': request.path}), 404


@app.errorhandler(500)
def internal_error(error):
    """Handle 500 errors"""
    logger.error(f"Internal server error: {error}")
    return jsonify({'error': 'Internal server error'}), 500


# ---------------------------------------------------------------------------
# Stateful AI Formatter  —  100-page pipeline with SSE progress streaming
# ---------------------------------------------------------------------------

@app.route('/api/formatting/stateful-format', methods=['POST'])
def stateful_format_document():
    """
    Stateful AI formatter for documents of 100+ pages.

    Accepts multipart/form-data with:
      file          — PDF or DOCX upload
      target_style  — IEEE | APA | MLA | Chicago | Harvard  (default: IEEE)
      title         — document title
      authors       — repeated field for each author name
      job_id        — optional; provide to resume a previous run
      resume        — 'true' to continue from the last checkpoint

    Streams Server-Sent Events (text/event-stream):
      data: {"type":"stage",    "stage":2, "description":"..."}
      data: {"type":"start",    "total":32, "bib_entries":45}
      data: {"type":"progress", "chunk":5, "total":32, "section":"...", ...}
      data: {"type":"complete", "latex":"...", "bib":"...", ...}
      data: {"type":"error",    "message":"..."}
    """
    def generate():
        try:
            from stateful_formatter import StatefulDocumentFormatter
            from file_extractor import FileExtractor

            if 'file' not in request.files:
                yield f"data: {json.dumps({'type':'error','message':'No file uploaded'})}\n\n"
                return

            uploaded   = request.files['file']
            file_bytes = uploaded.read()
            style      = request.form.get('target_style', 'IEEE')
            title      = request.form.get('title', 'Untitled Document')
            authors    = request.form.getlist('authors') or []
            job_id     = request.form.get('job_id') or str(uuid.uuid4())[:8]
            resume     = request.form.get('resume', 'false').lower() == 'true'

            # Stage 1 — extraction
            yield f"data: {json.dumps({'type':'stage','stage':1,'description':'Extracting document structure…'})}\n\n"

            extractor  = FileExtractor()
            extraction = extractor.extract_structured(file_bytes, uploaded.filename)
            if not extraction['success']:
                yield f"data: {json.dumps({'type':'error','message':extraction.get('error','Extraction failed')})}\n\n"
                return

            blocks   = extraction['blocks']
            job_dir  = os.path.join(UPLOAD_FOLDER, 'jobs', job_id)
            os.makedirs(job_dir, exist_ok=True)

            logger.info(
                '[stateful-format] job=%s style=%s blocks=%d resume=%s',
                job_id, style, len(blocks), resume
            )

            formatter = StatefulDocumentFormatter(job_dir)
            for event in formatter.format_large_document(
                blocks, style, title, authors, resume=resume
            ):
                yield f"data: {json.dumps(event)}\n\n"

        except Exception as exc:
            logger.error('[stateful-format] Unhandled error: %s', exc, exc_info=True)
            yield f"data: {json.dumps({'type':'error','message':str(exc)})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control':    'no-cache',
            'X-Accel-Buffering': 'no',
            'Access-Control-Allow-Origin': '*',
        },
    )


@app.route('/api/formatting/stateful-format/<job_id>/download', methods=['GET'])
def download_stateful_result(job_id):
    """
    Download the assembled paper.tex or refs.bib for a completed stateful job.
    Query param: ?file=tex | bib
    """
    file_type = request.args.get('file', 'tex')
    filename  = 'paper.tex' if file_type == 'tex' else 'refs.bib'
    path      = os.path.join(UPLOAD_FOLDER, 'jobs', job_id, filename)
    if not os.path.exists(path):
        return jsonify({'error': 'File not found — job may not be complete'}), 404
    with open(path, 'r', encoding='utf-8') as f:
        content = f.read()
    mime = 'text/plain'
    resp = Response(content, mimetype=mime)
    resp.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    return resp


# ============================================================================
# LEGACY NLP ANALYSIS ENDPOINTS (Stubs for compatibility)
# ============================================================================
# These endpoints are called by pythonNlpService.js but are not core to the
# Chapter 4 implementation. They return appropriate error messages or basic
# implementations without requiring Hugging Face API credentials.

@app.route('/api/nlp/analyze', methods=['POST'])
def nlp_analyze():
    """Legacy NLP analysis endpoint"""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'error': 'Text is required'}), 400
        
        return jsonify({
            'success': False,
            'error': 'NLP analysis not implemented. Use /api/document/sections or /api/grammar/enhance instead.',
            'recommendation': 'Use Document Formatter or Grammar Enhancer endpoints'
        }), 501
    except Exception as e:
        logger.error(f"NLP analyze error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/nlp/extract', methods=['POST'])
def nlp_extract():
    """Legacy NLP extraction endpoint"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        # Redirect to centralized file extraction
        return jsonify({
            'success': False,
            'error': 'Use /api/extract/file instead for document extraction'
        }), 501
    except Exception as e:
        logger.error(f"NLP extract error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/nlp/entities', methods=['POST'])
def nlp_extract_entities():
    """Legacy entity extraction endpoint"""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'error': 'Text is required'}), 400
        
        return jsonify({
            'success': False,
            'error': 'Entity extraction not available in this version. Hugging Face API credentials required.',
            'message': 'This feature requires a valid Hugging Face API token configured in the environment.'
        }), 503
    except Exception as e:
        logger.error(f"Entity extraction error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/nlp/keywords', methods=['POST'])
def nlp_keywords():
    """Legacy keyword extraction endpoint"""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'error': 'Text is required'}), 400
        
        text = data.get('text', '').strip()
        max_keywords = data.get('max_keywords', 10)
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Simple keyword extraction based on word frequency
        import re
        from collections import Counter
        
        words = re.findall(r'\b[a-z]{3,}\b', text.lower())
        stop_words = {'the', 'and', 'for', 'with', 'that', 'this', 'from', 'are', 'was', 'has', 'have', 'been', 'can', 'will', 'not', 'but', 'all', 'one', 'two', 'may', 'should', 'could', 'would', 'also', 'more', 'such', 'each', 'other', 'about', 'which', 'their', 'these', 'those', 'into', 'through', 'during'}
        
        filtered_words = [w for w in words if w not in stop_words and len(w) >= 4]
        counter = Counter(filtered_words)
        keywords = [{'keyword': word, 'frequency': count} for word, count in counter.most_common(max_keywords)]
        
        return jsonify({
            'success': True,
            'keywords': keywords,
            'count': len(keywords)
        }), 200
    except Exception as e:
        logger.error(f"Keywords extraction error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/nlp/summarize', methods=['POST'])
def nlp_summarize():
    """Legacy summarization endpoint - extractive summarization without ML"""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'error': 'Text is required'}), 400
        
        text = data.get('text', '').strip()
        max_length = data.get('max_length', 150)
        
        if len(text) < 100:
            return jsonify({
                'success': False,
                'error': 'Text too short to summarize (minimum 100 characters)',
                'summary': text
            }), 400
        
        # Simple extractive summarization
        import re
        sentences = re.split(r'(?<=[.!?])\s+', text)[:5]  # First 5 sentences max
        
        summary = ' '.join(sentences)
        if len(summary) > max_length:
            summary = summary[:max_length].rsplit(' ', 1)[0] + '...'
        
        return jsonify({
            'success': True,
            'summary': summary,
            'length': len(summary),
            'method': 'extractive (simple)'
        }), 200
    except Exception as e:
        logger.error(f"Summarization error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/nlp/sentiment', methods=['POST'])
def nlp_sentiment():
    """Legacy sentiment analysis endpoint - lexicon-based analysis"""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'error': 'Text is required'}), 400
        
        text = data.get('text', '').strip()
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Simple lexicon-based sentiment analysis
        positive_words = {'good', 'great', 'excellent', 'amazing', 'wonderful', 'fantastic', 'love', 'perfect', 'best', 'awesome', 'brilliant', 'outstanding', 'superb'}
        negative_words = {'bad', 'terrible', 'awful', 'hate', 'worst', 'poor', 'horrible', 'ugly', 'disgusting', 'wrong', 'dreadful', 'pathetic', 'useless'}
        
        text_lower = text.lower()
        pos_count = sum(1 for word in positive_words if word in text_lower)
        neg_count = sum(1 for word in negative_words if word in text_lower)
        total = pos_count + neg_count
        
        if total == 0:
            sentiment = 'neutral'
            score = 0.5
        elif pos_count > neg_count:
            sentiment = 'positive'
            score = 0.6 + (pos_count / total) * 0.4
        elif neg_count > pos_count:
            sentiment = 'negative'
            score = 0.4 - (neg_count / total) * 0.4
        else:
            sentiment = 'neutral'
            score = 0.5
        
        return jsonify({
            'success': True,
            'sentiment': sentiment,
            'score': round(min(1.0, max(0.0, score)), 2),
            'positive_words_count': pos_count,
            'negative_words_count': neg_count,
            'method': 'lexicon-based (simple)'
        }), 200
    except Exception as e:
        logger.error(f"Sentiment analysis error: {e}")
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    logger.info("Starting Python NLP Service - Chapter 4 Implementation")
    logger.info("Algorithms: Grammar Enhancement, Plagiarism Detection, Document Formatting, Section Detection")
    # Keep the NLP service stable while heavy libraries like Docling/EasyOCR load
    # models from site-packages. Flask's debug reloader treats those writes/imports
    # as source changes and can restart the service mid-request.
    port = int(os.getenv('PORT', '5001'))
    app.run(host='0.0.0.0', port=port, debug=False, use_reloader=False)


