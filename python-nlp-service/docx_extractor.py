"""
DOCX extraction with support for text, structure, tables, and embedded images
"""
from docx import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import base64
from io import BytesIO
import logging

logger = logging.getLogger(__name__)


def extract_docx_complete(filepath):
    """
    Extract text, structure, and media from DOCX
    
    Returns:
        dict: {
            'text': str,
            'word_count': int,
            'structure': list of elements,
            'media': list of media files
        }
    """
    extraction = {
        'text': '',
        'word_count': 0,
        'structure': [],
        'media': []
    }
    
    try:
        docx = DocxDocument(filepath)
        
        # ✅ Extract all elements preserving structure
        for element in docx.element.body:
            try:
                if isinstance(element, CT_P):  # Paragraph
                    para = None
                    for p in docx.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    
                    if para and para.text.strip():
                        text = para.text
                        extraction['text'] += text + '\n'
                        
                        # Determine if it's a heading
                        style_name = para.style.name if para.style else ''
                        if style_name.startswith('Heading'):
                            # Extract heading level from style name
                            try:
                                level = int(style_name.split()[-1]) if style_name[-1].isdigit() else 1
                            except:
                                level = 1
                            
                            extraction['structure'].append({
                                'type': 'heading',
                                'level': level,
                                'content': text
                            })
                        else:
                            extraction['structure'].append({
                                'type': 'paragraph',
                                'content': text
                            })
                
                elif isinstance(element, CT_Tbl):  # Table
                    table = None
                    for t in docx.tables:
                        if t._element == element:
                            table = t
                            break
                    
                    if table:
                        table_data = []
                        for row in table.rows:
                            table_data.append([cell.text for cell in row.cells])
                        
                        table_num = len([s for s in extraction['structure'] if s['type'] == 'table']) + 1
                        extraction['structure'].append({
                            'type': 'table',
                            'caption': f'Table {table_num}',
                            'rows': len(table.rows),
                            'columns': len(table.columns),
                            'data': table_data
                        })
                        
                        # Add table text to extraction
                        extraction['text'] += '\n'.join([' '.join(row) for row in table_data]) + '\n'
            except Exception as elem_err:
                logger.warning(f'Error processing document element: {elem_err}')
                continue
        
        # ✅ Extract embedded images - don't fail if this fails
        try:
            extraction['media'] = extract_images_from_docx(docx)
        except Exception as img_err:
            logger.warning(f'Error extracting images from DOCX: {img_err}')
            extraction['media'] = []
        
        extraction['word_count'] = len(extraction['text'].split())
        
    except Exception as e:
        logger.error(f'Error extracting DOCX: {e}')
        # Return partial extraction instead of raising
        if extraction['text']:  # If we got some text, return it
            extraction['word_count'] = len(extraction['text'].split())
            return extraction
        raise
    
    return extraction


def extract_images_from_docx(docx):
    """
    Extract embedded images from DOCX
    
    Args:
        docx: python-docx Document object
        
    Returns:
        list: List of image metadata dicts
    """
    images = []
    
    try:
        # Access all image relationships in the document
        for rel in docx.part.rels.values():
            if 'image' in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_blob = image_part.blob
                    
                    # Get image format from content type
                    content_type = image_part.content_type
                    image_ext = content_type.split('/')[-1]
                    
                    # Handle common formats
                    if image_ext == 'jpeg':
                        image_ext = 'jpg'
                    
                    filename = f'image_{len(images)}.{image_ext}'
                    
                    images.append({
                        'type': 'image',
                        'filename': filename,
                        'caption': f'Figure {len(images) + 1}',
                        'data': base64.b64encode(image_blob).decode(),
                        'encoding': 'base64',
                        'contentType': content_type
                    })
                    
                except Exception as e:
                    logger.warning(f'Error extracting individual image: {e}')
                    continue
    
    except Exception as e:
        logger.warning(f'Error accessing images from DOCX: {e}')
    
    return images


def iter_block_items(parent):
    """
    Generator for block-level elements (paragraphs and tables)
    Helps maintain document structure
    """
    if isinstance(parent, _Cell):
        parent_element = parent._tc
    else:
        parent_element = parent._element
    
    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
