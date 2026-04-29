"""Extraction and IR-building helpers."""

from __future__ import annotations

import json
import os
import re
import statistics
import subprocess
import sys
import tempfile
import time
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Tuple

import pdfplumber

from .config import REPO_ROOT, ensure_runtime_dirs, get_env


REFERENCE_HEADINGS = {"references", "bibliography", "works cited"}
ABSTRACT_HEADINGS = {"abstract"}
KEYWORD_PATTERN = re.compile(r"^keywords?\s*[:.-]\s*(.+)$", re.IGNORECASE)
COMMON_SECTION_HEADINGS = {
    "abstract",
    "introduction",
    "background",
    "related work",
    "literature review",
    "method",
    "methods",
    "methodology",
    "materials and methods",
    "results",
    "discussion",
    "conclusion",
    "conclusions",
    "acknowledgment",
    "acknowledgments",
    "references",
}
FRONTMATTER_HEADINGS = {
    "table of contents",
    "contents",
    "list of tables",
    "list of figures",
}


def extract_to_ir(input_path: str) -> Tuple[Dict[str, Any], List[str]]:
    source_path = Path(input_path).expanduser().resolve()
    if not source_path.exists():
        raise FileNotFoundError(f"Input file not found: {source_path}")

    warnings: List[str] = []
    runtime_dirs = ensure_runtime_dirs()
    suffix = source_path.suffix.lower()

    if suffix == ".pdf":
        extractor_mode = str(get_env("FORMATTING_ENGINE_EXTRACTOR") or "auto").strip().lower()
        warnings.append("Using native pdfplumber/Nougat/Docling-capable structure extraction.")
        native_ir = _build_ir_from_pdf_native(source_path, warnings)
        ir = native_ir
        if extractor_mode in {"auto", "docling", "docling-auto"}:
            docling_ir, docling_warnings = _build_ir_from_docling(source_path)
            warnings.extend(docling_warnings)
            if docling_ir:
                native_score = _ir_quality_score(ir)
                docling_score = _ir_quality_score(docling_ir)
                if extractor_mode == "docling" or docling_score > native_score:
                    warnings.append(
                        f"Docling extraction selected (score {docling_score:.1f} vs current {native_score:.1f})."
                    )
                    ir = docling_ir
                else:
                    warnings.append(
                        f"Current PDF extraction kept over Docling (score {native_score:.1f} vs docling {docling_score:.1f})."
                    )
            elif extractor_mode == "docling":
                warnings.append("Docling was requested but unavailable/failed; used native PDF extraction.")
        if extractor_mode in {"auto", "nougat", "nougat-ocr"}:
            nougat_ir, nougat_warnings = _build_ir_from_pdf_nougat(source_path)
            warnings.extend(nougat_warnings)
            if nougat_ir:
                native_score = _ir_quality_score(ir)
                nougat_score = _ir_quality_score(nougat_ir)
                if extractor_mode in {"nougat", "nougat-ocr"} or nougat_score > native_score:
                    warnings.append(
                        f"Nougat OCR extraction selected (score {nougat_score:.1f} vs current {native_score:.1f})."
                    )
                    ir = nougat_ir
                else:
                    warnings.append(
                        f"Current PDF extraction kept over Nougat OCR (score {native_score:.1f} vs nougat {nougat_score:.1f})."
                    )
            elif extractor_mode in {"nougat", "nougat-ocr"}:
                warnings.append("Nougat OCR was requested but unavailable/failed; used native PDF extraction.")
        ir["figures"] = _extract_figures_with_pymupdf(source_path, runtime_dirs["images"], warnings)
        ir["tables"] = _extract_tables_with_pdfplumber(source_path, warnings)
        recovered_references = _extract_pdf_references_from_text(source_path, warnings)
        if len(recovered_references) > len(ir.get("references") or []):
            ir["references"] = recovered_references
            warnings.append(f"Recovered {len(recovered_references)} references from raw PDF text.")
    elif suffix == ".docx":
        extractor_mode = str(get_env("FORMATTING_ENGINE_EXTRACTOR") or "auto").strip().lower()
        warnings.append("Using native DOCX/Docling-capable structure extraction.")
        native_ir = _build_ir_from_docx_native(source_path)
        ir = native_ir
        if extractor_mode in {"auto", "docling", "docling-auto"}:
            docling_ir, docling_warnings = _build_ir_from_docling(source_path)
            warnings.extend(docling_warnings)
            if docling_ir:
                native_score = _ir_quality_score(native_ir)
                docling_score = _ir_quality_score(docling_ir)
                if extractor_mode == "docling" or docling_score > native_score:
                    warnings.append(
                        f"Docling extraction selected (score {docling_score:.1f} vs native {native_score:.1f})."
                    )
                    ir = docling_ir
                else:
                    warnings.append(
                        f"Native DOCX extraction kept over Docling (score {native_score:.1f} vs docling {docling_score:.1f})."
                    )
            elif extractor_mode == "docling":
                warnings.append("Docling was requested but unavailable/failed; used native DOCX extraction.")
    else:
        markdown = _fallback_text_extraction(source_path)
        ir = _build_ir_from_markdown(markdown, source_path)

    _attach_semantic_ir_metadata(ir)
    ir.setdefault("debug", {}).setdefault("docling", {"enabled": False, "reason": "native-fallback"})
    return ir, warnings


def _build_ir_from_docx_native(source_path: Path) -> Dict[str, Any]:
    return _blocks_to_ir_with_optional_chunker(_extract_docx_native_blocks(source_path), source_path)


def _build_ir_from_pdf_native(source_path: Path, warnings: List[str]) -> Dict[str, Any]:
    try:
        blocks = _extract_pdf_native_blocks(source_path, warnings)
    except Exception as exc:
        warnings.append(f"Native PDF structure extraction failed: {exc}; using plain text fallback.")
        return _build_ir_from_markdown(_fallback_text_extraction(source_path), source_path)

    return _blocks_to_ir_with_optional_chunker(blocks, source_path)


def _build_ir_from_docling(source_path: Path) -> Tuple[Dict[str, Any] | None, List[str]]:
    warnings: List[str] = []
    try:
        from docling.document_converter import DocumentConverter  # type: ignore
    except Exception as exc:
        return None, [f"Docling is not installed or failed to import; skipped Docling extraction: {exc}"]

    try:
        result = DocumentConverter().convert(str(source_path))
        document = getattr(result, "document", None)
        if document is None:
            return None, ["Docling completed but returned no document."]
        markdown = ""
        if hasattr(document, "export_to_markdown"):
            markdown = document.export_to_markdown()
        elif hasattr(document, "export_to_text"):
            markdown = document.export_to_text()
        elif hasattr(document, "model_dump"):
            dumped = document.model_dump()
            if isinstance(dumped, dict):
                markdown = str(dumped.get("body") or "")
        markdown = str(markdown or "").strip()
        if len(markdown) < 200:
            return None, ["Docling output was too short to use."]
        ir = _build_ir_from_markdown(markdown, source_path)
        ir.setdefault("source", {})["ir_builder"] = "docling"
        ir.setdefault("debug", {})["docling"] = {
            "enabled": True,
            "markdown_chars": len(markdown),
        }
        return ir, warnings
    except Exception as exc:
        return None, [f"Docling extraction failed: {exc}"]


def _build_ir_from_pdf_nougat(source_path: Path) -> Tuple[Dict[str, Any] | None, List[str]]:
    warnings: List[str] = []
    nougat_command = _nougat_command()
    if not nougat_command:
        return None, ["Nougat OCR is not installed or not on PATH; skipped Nougat extraction."]

    timeout_seconds = _int_env("NOUGAT_TIMEOUT_SECONDS", 900)
    model = str(get_env("NOUGAT_MODEL") or "").strip()
    checkpoint = str(get_env("NOUGAT_CHECKPOINT") or "").strip()
    pages = str(get_env("NOUGAT_PAGES") or "").strip()
    no_skipping = _truthy_env("NOUGAT_NO_SKIPPING")

    runtime_root = ensure_runtime_dirs()["runtime"]
    temp_parent = runtime_root / "tmp"
    temp_parent.mkdir(parents=True, exist_ok=True)

    temp_dir = tempfile.mkdtemp(prefix="documentor_nougat_", dir=str(temp_parent))
    command = nougat_command + [str(source_path), "-o", temp_dir, "--markdown"]
    if checkpoint:
        command.extend(["--checkpoint", checkpoint])
    if model:
        command.extend(["--model", model])
    if pages:
        command.extend(["--pages", pages])
    if no_skipping:
        command.append("--no-skipping")

    try:
        env = os.environ.copy()
        cache_dir = ensure_runtime_dirs()["runtime"] / "nougat_cache"
        cache_dir.mkdir(parents=True, exist_ok=True)
        env.setdefault("XDG_CACHE_HOME", str(cache_dir))
        env.setdefault("TORCH_HOME", str(cache_dir / "torch"))
        env.setdefault("HF_HOME", str(cache_dir / "huggingface"))
        env.setdefault("NO_ALBUMENTATIONS_UPDATE", "1")
        env.setdefault("TF_CPP_MIN_LOG_LEVEL", "2")
        completed = subprocess.run(
            command,
            capture_output=True,
            text=True,
            check=False,
            timeout=timeout_seconds,
            env=env,
        )
    except subprocess.TimeoutExpired:
        return None, [f"Nougat OCR timed out after {timeout_seconds}s; skipped Nougat extraction."]
    except Exception as exc:
        return None, [f"Nougat OCR failed to start: {exc}"]

    if completed.returncode != 0:
        stderr = _trim_warning(completed.stderr or completed.stdout)
        return None, [f"Nougat OCR exited with code {completed.returncode}: {stderr}"]

    mmd_files = sorted(Path(temp_dir).rglob("*.mmd"), key=lambda path: path.stat().st_mtime, reverse=True)
    if not mmd_files:
        text_files = sorted(Path(temp_dir).rglob("*.md"), key=lambda path: path.stat().st_mtime, reverse=True)
        mmd_files = text_files
    if not mmd_files:
        return None, ["Nougat OCR completed but produced no Markdown/MMD output."]

    markdown = mmd_files[0].read_text(encoding="utf-8", errors="replace")
    missing_pages = markdown.count("[MISSING_PAGE]")
    if missing_pages:
        warnings.append(f"Nougat OCR output contains {missing_pages} [MISSING_PAGE] marker(s).")

    cleaned_markdown = markdown.replace("[MISSING_PAGE]", "\n")
    if len(cleaned_markdown.strip()) < 500:
        return None, ["Nougat OCR output was too short to use."]

    ir = _build_ir_from_markdown(cleaned_markdown, source_path)
    ir.setdefault("source", {})["ir_builder"] = "nougat_ocr"
    ir.setdefault("debug", {})["nougat"] = {
        "enabled": True,
        "model": model or "default",
        "checkpoint": checkpoint or None,
        "pages": pages or "all",
        "missing_pages": missing_pages,
        "markdown_chars": len(cleaned_markdown),
    }
    return ir, warnings


def compare_chunker_ir(input_path: str) -> Dict[str, Any]:
    """Compare legacy IR building against the shared DocumentChunker path."""
    source_path = Path(input_path).expanduser().resolve()
    if not source_path.exists():
        raise FileNotFoundError(f"Input file not found: {source_path}")

    warnings: List[str] = []
    suffix = source_path.suffix.lower()
    if suffix == ".pdf":
        blocks = _extract_pdf_native_blocks(source_path, warnings)
    elif suffix == ".docx":
        blocks = _extract_docx_native_blocks(source_path)
    else:
        markdown = _fallback_text_extraction(source_path)
        blocks = _markdown_to_blocks(markdown)

    legacy_started = time.perf_counter()
    legacy_ir = _blocks_to_ir(blocks, source_path)
    legacy_ms = round((time.perf_counter() - legacy_started) * 1000, 2)

    chunker_started = time.perf_counter()
    chunker_ir, chunker_warnings = _blocks_to_ir_with_document_chunker(blocks, source_path)
    chunker_ms = round((time.perf_counter() - chunker_started) * 1000, 2)

    return {
        "input_path": str(source_path),
        "block_count": len(blocks),
        "warnings": warnings + chunker_warnings,
        "legacy": _ir_metrics(legacy_ir, legacy_ms),
        "document_chunker": _ir_metrics(chunker_ir, chunker_ms),
        "deltas": _ir_metric_deltas(legacy_ir, chunker_ir, legacy_ms, chunker_ms),
    }


def _extract_docx_native_blocks(source_path: Path) -> List[Dict[str, Any]]:
    try:
        from docx import Document as WordDocument  # type: ignore
    except Exception as exc:
        raise RuntimeError("DOCX fallback requires python-docx") from exc

    document = WordDocument(str(source_path))
    blocks: List[Dict[str, Any]] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "Normal").lower()
        heading_match = re.search(r"heading\s*(\d+)", style_name)
        source_format = _source_format_from_docx_paragraph(paragraph)
        if style_name == "title":
            blocks.append({"type": "title", "level": 1, "text": text, "source_format": source_format})
        elif heading_match:
            blocks.append({"type": "heading", "level": int(heading_match.group(1)), "text": text, "source_format": source_format})
        elif _looks_like_heading_text(text):
            blocks.append({"type": "heading", "level": 1, "text": _clean_heading_text(text), "source_format": source_format})
        else:
            blocks.append({"type": "paragraph", "text": text, "source_format": source_format})

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            blocks.append({"type": "table", "id": f"table-{table_index}", "rows": rows})

    return blocks


def _extract_pdf_native_blocks(source_path: Path, warnings: List[str]) -> List[Dict[str, Any]]:
    blocks: List[Dict[str, Any]] = []
    with pdfplumber.open(source_path) as pdf:
        pages = pdf.pages
        if not pages:
            raise ValueError("PDF is empty")
        raw_lines_by_page = [_extract_pdf_lines(page, page_index + 1) for page_index, page in enumerate(pages)]
        noise = _detect_repeating_pdf_noise(raw_lines_by_page, pages)
        _annotate_pdf_line_zones(raw_lines_by_page, noise)
        body_sizes = [
            line["size"]
            for lines in raw_lines_by_page
            for line in lines
            if line["text"]
            and line["fingerprint"] not in noise
            and not _is_page_number(line["text"])
            and line.get("zone") == "main_body"
        ]
        body_size = statistics.median(body_sizes) if body_sizes else 10.0

        paragraph_lines: List[Dict[str, Any]] = []

        def flush_paragraph() -> None:
            nonlocal paragraph_lines
            block = _pdf_lines_to_paragraph_block(paragraph_lines)
            if block:
                blocks.append(block)
            paragraph_lines = []

        for page, lines in zip(pages, raw_lines_by_page):
            if _is_frontmatter_page(lines):
                flush_paragraph()
                continue

            previous_line: Dict[str, Any] | None = None
            for line in lines:
                text = line["text"].strip()
                if not text or line["fingerprint"] in noise or _is_page_number(text):
                    continue
                if _should_skip_pdf_body_line(line, body_size):
                    continue

                is_heading = _is_pdf_heading(line, body_size)
                if is_heading:
                    flush_paragraph()
                    blocks.append(
                        {
                            "type": "heading",
                            "level": _heading_level_from_text(text),
                            "text": _clean_heading_text(text),
                            "source_format": _source_format_from_pdf_lines([line]),
                        }
                    )
                    previous_line = line
                    continue

                if previous_line:
                    gap = line["top"] - previous_line["bottom"]
                    indent_delta = abs(line["x0"] - previous_line["x0"])
                    if gap > body_size * 0.75 or indent_delta > 24:
                        flush_paragraph()

                paragraph_lines.append(line)
                previous_line = line

        flush_paragraph()

    return blocks


def _extract_pdf_lines(page: Any, page_number: int) -> List[Dict[str, Any]]:
    words = page.extract_words(
        extra_attrs=["size", "fontname"],
        use_text_flow=False,
        keep_blank_chars=False,
        x_tolerance=3,
        y_tolerance=3,
    ) or []
    if not words:
        return []

    words = sorted(words, key=lambda item: (round(float(item.get("top", 0)), 1), float(item.get("x0", 0))))
    lines: List[List[Dict[str, Any]]] = []
    for word in words:
        if lines and abs(float(word.get("top", 0)) - float(lines[-1][0].get("top", 0))) <= 3:
            lines[-1].append(word)
        else:
            lines.append([word])

    result = []
    for line_words in lines:
        line_words = sorted(line_words, key=lambda item: float(item.get("x0", 0)))
        for line_words in _split_line_by_horizontal_gaps(line_words):
            if not line_words:
                continue
            line = _pdf_word_group_to_line(line_words, page, page_number)
            if line:
                result.append(line)
    return result


def _split_line_by_horizontal_gaps(line_words: List[Dict[str, Any]]) -> List[List[Dict[str, Any]]]:
    if len(line_words) <= 1:
        return [line_words]

    groups: List[List[Dict[str, Any]]] = [[line_words[0]]]
    for previous, current in zip(line_words, line_words[1:]):
        gap = float(current.get("x0", 0)) - float(previous.get("x1", 0))
        if gap > 24:
            groups.append([current])
        else:
            groups[-1].append(current)
    return groups


def _pdf_word_group_to_line(line_words: List[Dict[str, Any]], page: Any, page_number: int) -> Dict[str, Any]:
        text = " ".join(str(word.get("text", "")) for word in line_words).strip()
        if not text:
            return {}
        sizes = [float(word.get("size") or 0) for word in line_words if word.get("size")]
        fonts = [str(word.get("fontname") or "") for word in line_words]
        return {
            "page": page_number,
            "text": text,
            "fingerprint": _line_fingerprint(text),
            "top": min(float(word.get("top", 0)) for word in line_words),
            "bottom": max(float(word.get("bottom", 0)) for word in line_words),
            "x0": min(float(word.get("x0", 0)) for word in line_words),
            "x1": max(float(word.get("x1", 0)) for word in line_words),
            "size": statistics.mean(sizes) if sizes else 0,
            "fontname": Counter(fonts).most_common(1)[0][0] if fonts else "",
            "bold": any("bold" in font.lower() for font in fonts),
            "italic": any("italic" in font.lower() or "oblique" in font.lower() for font in fonts),
            "page_width": float(page.width),
            "page_height": float(page.height),
        }


def _pdf_lines_to_paragraph_block(lines: List[Dict[str, Any]]) -> Dict[str, Any] | None:
    if not lines:
        return None
    text = " ".join(str(line.get("text") or "").strip() for line in lines).strip()
    text = re.sub(r"\s+", " ", text)
    text = _clean_pdf_paragraph_text(text)
    if not text:
        return None
    return {
        "type": "paragraph",
        "text": text,
        "source_format": _source_format_from_pdf_lines(lines),
    }


def _source_format_from_pdf_lines(lines: List[Dict[str, Any]]) -> Dict[str, Any]:
    if not lines:
        return {"source": "pdf"}

    sizes = [float(line.get("size") or 0) for line in lines if line.get("size")]
    fonts = [str(line.get("fontname") or "") for line in lines if line.get("fontname")]
    page_widths = [float(line.get("page_width") or 0) for line in lines if line.get("page_width")]
    lefts = [float(line.get("x0") or 0) for line in lines]
    rights = [float(line.get("x1") or 0) for line in lines]
    tops = [float(line.get("top") or 0) for line in lines]
    bottoms = [float(line.get("bottom") or 0) for line in lines]
    bold_count = sum(1 for line in lines if line.get("bold"))
    italic_count = sum(1 for line in lines if line.get("italic"))
    page_width = statistics.median(page_widths) if page_widths else 612.0
    x0 = statistics.median(lefts) if lefts else 0.0
    x1 = statistics.median(rights) if rights else page_width
    zones = [str(line.get("zone") or "main_body") for line in lines]
    zone = Counter(zones).most_common(1)[0][0] if zones else "main_body"

    return {
        "source": "pdf",
        "zone": zone,
        "semantic_role": _semantic_role_from_zone(zone),
        "font_name": _clean_pdf_font_name(Counter(fonts).most_common(1)[0][0]) if fonts else None,
        "font_size_pt": round(statistics.median(sizes), 2) if sizes else None,
        "bold": bold_count >= max(1, len(lines) / 2),
        "italic": italic_count >= max(1, len(lines) / 2),
        "alignment": _infer_pdf_alignment(x0, x1, page_width),
        "left_indent_inches": round(max(0.0, x0) / 72.0, 2),
        "right_indent_inches": round(max(0.0, page_width - x1) / 72.0, 2),
        "line_count": len(lines),
        "bbox": {
            "x0": round(min(lefts), 2) if lefts else 0,
            "x1": round(max(rights), 2) if rights else 0,
            "top": round(min(tops), 2) if tops else 0,
            "bottom": round(max(bottoms), 2) if bottoms else 0,
        },
    }


def _clean_pdf_paragraph_text(text: str) -> str:
    value = str(text or "").strip()
    cleanup_patterns = [
        r"^Submitted for possible open access\s+",
        r"^Universal Journal of Computer Sciences and Communications\s+",
        r"^Copyright:\s*©?\s*\d{4}\s+by the authors\.\s*",
        r"^Attribution \(CC BY\) license\s+",
    ]
    for pattern in cleanup_patterns:
        value = re.sub(pattern, "", value, flags=re.IGNORECASE).strip()
    return value


def _source_format_from_docx_paragraph(paragraph: Any) -> Dict[str, Any]:
    paragraph_format = paragraph.paragraph_format
    runs = [run for run in paragraph.runs if run.text]
    font_names = [run.font.name for run in runs if run.font.name]
    sizes = [run.font.size.pt for run in runs if run.font.size]
    bold_values = [bool(run.bold) for run in runs if run.bold is not None]
    italic_values = [bool(run.italic) for run in runs if run.italic is not None]

    return {
        "source": "docx",
        "zone": "main_body",
        "semantic_role": "body",
        "style_name": paragraph.style.name if paragraph.style else None,
        "font_name": Counter(font_names).most_common(1)[0][0] if font_names else None,
        "font_size_pt": round(statistics.median(sizes), 2) if sizes else None,
        "bold": bool(bold_values and sum(1 for value in bold_values if value) >= len(bold_values) / 2),
        "italic": bool(italic_values and sum(1 for value in italic_values if value) >= len(italic_values) / 2),
        "alignment": _docx_alignment_name(paragraph.alignment),
        "left_indent_inches": _length_inches(paragraph_format.left_indent),
        "first_line_indent_inches": _length_inches(paragraph_format.first_line_indent),
        "space_before_pt": _length_points(paragraph_format.space_before),
        "space_after_pt": _length_points(paragraph_format.space_after),
        "line_spacing": float(paragraph_format.line_spacing) if isinstance(paragraph_format.line_spacing, (int, float)) else None,
    }


def _clean_pdf_font_name(fontname: str) -> str:
    value = str(fontname or "")
    if "+" in value:
        value = value.split("+", 1)[1]
    value = re.sub(r"[-,]?(?:Roman|Regular|BoldItalic|Bold|Italic|MT|PS)$", "", value, flags=re.IGNORECASE)
    value = value.replace("Linotype", "").replace("PS", "")
    value = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", value)
    return re.sub(r"\s+", " ", value).strip() or fontname


def _infer_pdf_alignment(x0: float, x1: float, page_width: float) -> str:
    left_margin = x0
    right_margin = page_width - x1
    center_delta = abs((x0 + x1) / 2 - page_width / 2)
    if center_delta < 18 and abs(left_margin - right_margin) < 36:
        return "center"
    if right_margin < 36 and left_margin > 72:
        return "right"
    return "left"


def _docx_alignment_name(alignment: Any) -> str | None:
    if alignment is None:
        return None
    name = getattr(alignment, "name", None)
    if name:
        return str(name).lower()
    return str(alignment).lower()


def _length_inches(value: Any) -> float | None:
    if value is None:
        return None
    try:
        return round(float(value.inches), 2)
    except Exception:
        return None


def _length_points(value: Any) -> float | None:
    if value is None:
        return None
    try:
        return round(float(value.pt), 2)
    except Exception:
        return None


def _blocks_to_ir_with_optional_chunker(blocks: List[Dict[str, Any]], source_path: Path) -> Dict[str, Any]:
    mode = str(get_env("FORMATTING_ENGINE_USE_DOCUMENT_CHUNKER", "auto")).strip().lower()
    if mode in {"0", "false", "no", "off", "legacy"}:
        ir = _blocks_to_ir(blocks, source_path)
        ir.setdefault("debug", {})["document_chunker"] = {"enabled": False, "used": False, "mode": mode}
        return ir

    legacy = _blocks_to_ir(blocks, source_path)
    if mode in {"auto", ""} and not _legacy_ir_needs_chunker(legacy, blocks):
        legacy.setdefault("debug", {})["document_chunker"] = {
            "enabled": True,
            "used": False,
            "mode": "auto",
            "reason": "legacy_ir_sufficient",
        }
        return legacy

    ir, warnings = _blocks_to_ir_with_document_chunker(blocks, source_path)
    if warnings:
        legacy.setdefault("debug", {})["document_chunker"] = {
            "enabled": True,
            "used": False,
            "mode": mode,
            "warnings": warnings,
        }
        return legacy

    if mode in {"auto", ""} and not _chunker_ir_is_better(legacy, ir):
        legacy.setdefault("debug", {})["document_chunker"] = {
            "enabled": True,
            "used": False,
            "mode": "auto",
            "reason": "chunker_ir_not_better",
        }
        return legacy

    ir.setdefault("debug", {})["document_chunker"] = {
        "enabled": True,
        "used": True,
        "mode": mode,
        "warnings": [],
    }
    return ir


def _legacy_ir_needs_chunker(ir: Dict[str, Any], blocks: List[Dict[str, Any]]) -> bool:
    sections = ir.get("sections") or []
    paragraphs = [paragraph for section in sections for paragraph in section.get("paragraphs", [])]
    heading_blocks = [block for block in blocks if block.get("type") in {"heading", "title"}]
    if len(sections) <= 1 and len(heading_blocks) >= 3:
        return True
    if not paragraphs and len(blocks) >= 20:
        return True
    if not (ir.get("metadata") or {}).get("abstract") and any(str(block.get("text") or "").strip().lower() == "abstract" for block in heading_blocks):
        return True
    return False


def _chunker_ir_is_better(legacy: Dict[str, Any], chunker: Dict[str, Any]) -> bool:
    legacy_sections = legacy.get("sections") or []
    chunker_sections = chunker.get("sections") or []
    legacy_paragraphs = [paragraph for section in legacy_sections for paragraph in section.get("paragraphs", [])]
    chunker_paragraphs = [paragraph for section in chunker_sections for paragraph in section.get("paragraphs", [])]

    score = 0
    if len(chunker_sections) > len(legacy_sections):
        score += 2
    if len(chunker_paragraphs) > len(legacy_paragraphs):
        score += 2
    if len(chunker.get("references") or []) > len(legacy.get("references") or []):
        score += 1
    if (chunker.get("metadata") or {}).get("abstract") and not (legacy.get("metadata") or {}).get("abstract"):
        score += 1
    if len(chunker_paragraphs) < len(legacy_paragraphs):
        score -= 3
    if len(chunker_sections) < len(legacy_sections):
        score -= 3

    return score > 0


def _blocks_to_ir_with_document_chunker(blocks: List[Dict[str, Any]], source_path: Path) -> Tuple[Dict[str, Any], List[str]]:
    warnings: List[str] = []
    try:
        service_path = str(REPO_ROOT / "python-nlp-service")
        if service_path not in sys.path:
            sys.path.insert(0, service_path)
        from document_chunker import DocumentChunker  # type: ignore
    except Exception as exc:
        return _blocks_to_ir(blocks, source_path), [f"DocumentChunker unavailable: {exc}"]

    structured_blocks = _blocks_to_structured_blocks(blocks)
    toc = [
        {
            "level": block.get("section_level", 1),
            "text": block.get("text", ""),
            "index": block.get("index", index),
        }
        for index, block in enumerate(structured_blocks)
        if int(block.get("section_level") or 0) > 0
    ]

    title = _first_block_text(blocks) or source_path.stem
    try:
        chunks = DocumentChunker().chunk_blocks(
            structured_blocks,
            title=title,
            table_of_contents=toc,
        )
    except Exception as exc:
        return _blocks_to_ir(blocks, source_path), [f"DocumentChunker failed: {exc}"]

    if not chunks:
        return _blocks_to_ir(blocks, source_path), ["DocumentChunker returned no chunks."]

    metadata = {
        "title": title,
        "authors": [],
        "affiliation": "",
        "abstract": "",
        "keywords": [],
        "date": None,
    }
    sections: List[Dict[str, Any]] = []
    references: List[str] = []
    table_blocks = [block for block in blocks if block.get("type") == "table"]

    for chunk in chunks:
        heading = str(chunk.get("heading") or "").strip()
        section_type = str(chunk.get("section_type") or "generic").strip().lower()
        text = str(chunk.get("text") or "").strip()
        heading_lower = heading.lower()

        if not text and not heading:
            continue
        if heading_lower == metadata["title"].lower() and section_type == "generic":
            continue
        if heading_lower in FRONTMATTER_HEADINGS:
            continue

        if section_type == "abstract" or heading_lower in ABSTRACT_HEADINGS:
            metadata["abstract"] = _strip_leading_heading(text, heading)
            continue

        if section_type == "references" or heading_lower in REFERENCE_HEADINGS:
            references.extend(_split_reference_lines(_strip_leading_heading(text, heading).splitlines()))
            continue

        paragraphs = _split_chunk_paragraphs(_strip_leading_heading(text, heading))
        if not paragraphs:
            continue

        sections.append(
            {
                "level": int(chunk.get("section_level") or 1),
                "title": _clean_heading_text(heading or f"Section {len(sections) + 1}"),
                "paragraphs": paragraphs,
                "paragraph_formats": [{} for _ in paragraphs],
                "figures": [],
                "tables": [],
            }
        )

    keyword_match = KEYWORD_PATTERN.search(metadata.get("abstract") or "")
    if keyword_match:
        metadata["keywords"] = [part.strip() for part in re.split(r"[;,]", keyword_match.group(1)) if part.strip()]

    return {
        "metadata": metadata,
        "sections": sections or [_new_section(1, "Introduction")],
        "figures": [],
        "tables": table_blocks,
        "references": references,
        "footnotes": [],
        "source": {
            "path": str(source_path),
            "suffix": source_path.suffix.lower(),
            "ir_builder": "document_chunker",
        },
    }, warnings


def _blocks_to_structured_blocks(blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    structured: List[Dict[str, Any]] = []
    for index, block in enumerate(blocks):
        block_type = block.get("type")
        text = str(block.get("text") or "").strip()
        if block_type == "table":
            rows = block.get("rows") or []
            text = "\n".join(" | ".join(str(cell or "").strip() for cell in row) for row in rows if row)
            style = "Normal"
            level = 0
        elif block_type == "title":
            style = "Title"
            level = 1
        elif block_type == "heading":
            level = int(block.get("level") or 1)
            style = f"Heading {max(1, min(level, 6))}"
        else:
            style = "Normal"
            level = 0

        if not text:
            continue
        structured.append(
            {
                "text": text,
                "style": style,
                "section_level": level,
                "index": index,
                "word_count": len(text.split()),
                "char_count": len(text),
            }
        )
    return structured


def _strip_leading_heading(text: str, heading: str) -> str:
    value = str(text or "").strip()
    heading = str(heading or "").strip()
    if heading and value.lower().startswith(heading.lower()):
        value = value[len(heading):].strip(" \n\r\t:-")
    return value


def _split_chunk_paragraphs(text: str) -> List[str]:
    paragraphs = [re.sub(r"\s+", " ", paragraph).strip() for paragraph in re.split(r"\n{2,}", text or "")]
    return [paragraph for paragraph in paragraphs if paragraph]


def _detect_repeating_pdf_noise(lines_by_page: List[List[Dict[str, Any]]], pages: List[Any]) -> set[str]:
    counts: Counter[str] = Counter()
    page_count = len(lines_by_page)
    if page_count < 3:
        return set()

    for page, lines in zip(pages, lines_by_page):
        top_limit = float(page.height) * 0.10
        bottom_limit = float(page.height) * 0.90
        seen_on_page = {
            line["fingerprint"]
            for line in lines
            if line["top"] <= top_limit or line["bottom"] >= bottom_limit
        }
        counts.update(seen_on_page)

    min_repeats = max(2, int(page_count * 0.35))
    return {fingerprint for fingerprint, count in counts.items() if fingerprint and count >= min_repeats}


def _annotate_pdf_line_zones(lines_by_page: List[List[Dict[str, Any]]], repeating_noise: set[str]) -> None:
    """Classify page regions before section detection consumes line text."""
    for lines in lines_by_page:
        for line in lines:
            line["zone"] = _pdf_line_zone(line, repeating_noise)


def _pdf_line_zone(line: Dict[str, Any], repeating_noise: set[str]) -> str:
    text = str(line.get("text") or "").strip()
    page_width = float(line.get("page_width") or 612.0)
    page_height = float(line.get("page_height") or 792.0)
    x0 = float(line.get("x0") or 0.0)
    x1 = float(line.get("x1") or 0.0)
    top = float(line.get("top") or 0.0)
    bottom = float(line.get("bottom") or 0.0)
    width_ratio = max(0.0, x1 - x0) / max(page_width, 1.0)
    center_ratio = ((x0 + x1) / 2.0) / max(page_width, 1.0)

    if line.get("fingerprint") in repeating_noise or top <= page_height * 0.07:
        return "header"
    if bottom >= page_height * 0.93 or _is_page_number(text):
        return "footer"
    if _looks_like_metadata_or_sidebar_line(text):
        return "marginalia"
    if width_ratio <= 0.26 and (center_ratio <= 0.18 or center_ratio >= 0.82):
        return "sidebar"
    return "main_body"


def _semantic_role_from_zone(zone: str) -> str:
    if zone in {"header", "footer"}:
        return "running_matter"
    if zone in {"sidebar", "marginalia"}:
        return "marginal_content"
    return "body"


def _attach_semantic_ir_metadata(ir: Dict[str, Any]) -> None:
    sections = ir.get("sections") or []
    zone_counts: Counter[str] = Counter()
    role_counts: Counter[str] = Counter()
    for section in sections:
        section.setdefault("semantic_role", "body_section")
        for source_format in section.get("paragraph_formats") or []:
            if not isinstance(source_format, dict):
                continue
            zone = str(source_format.get("zone") or "main_body")
            role = str(source_format.get("semantic_role") or _semantic_role_from_zone(zone))
            source_format["zone"] = zone
            source_format["semantic_role"] = role
            zone_counts[zone] += 1
            role_counts[role] += 1

    ir.setdefault("semantic_ir", {})["roles"] = {
        "metadata": "front_matter",
        "sections": "body",
        "references": "back_matter",
        "figures": "float",
        "tables": "float",
    }
    ir["semantic_ir"]["zone_counts"] = dict(zone_counts)
    ir["semantic_ir"]["role_counts"] = dict(role_counts)
    ir.setdefault("source", {})["zone_detector"] = {
        "enabled": True,
        "strategy": "pdf_geometry_heuristic",
    }


def _blocks_to_ir(blocks: List[Dict[str, Any]], source_path: Path) -> Dict[str, Any]:
    title = _first_block_text(blocks) or source_path.stem
    metadata = {
        "title": title,
        "authors": [],
        "affiliation": "",
        "abstract": "",
        "keywords": [],
        "date": None,
    }
    sections: List[Dict[str, Any]] = []
    references: List[str] = []
    current: Dict[str, Any] | None = None
    in_references = False
    in_abstract = False
    abstract_lines: List[str] = []
    reference_lines: List[str] = []

    def flush_current() -> None:
        nonlocal current
        if current is None:
            return
        filtered = [
            (paragraph, current["paragraph_formats"][index] if index < len(current.get("paragraph_formats", [])) else {})
            for index, paragraph in enumerate(current["paragraphs"])
            if paragraph.strip()
        ]
        current["paragraphs"] = [paragraph for paragraph, _ in filtered]
        current["paragraph_formats"] = [source_format for _, source_format in filtered]
        if current["title"] or current["paragraphs"] or current["tables"]:
            sections.append(current)
        current = None

    for block in blocks:
        text = str(block.get("text") or "").strip()
        block_type = block.get("type")
        if not text and block_type != "table":
            continue

        lower = text.lower()
        if block_type == "title":
            metadata["title"] = text
            continue

        if block_type == "heading":
            clean = _clean_heading_text(text)
            clean_lower = clean.lower()
            if clean_lower and clean_lower in str(metadata.get("title") or "").lower():
                continue
            if clean_lower in REFERENCE_HEADINGS:
                flush_current()
                in_references = True
                in_abstract = False
                continue
            if clean_lower in ABSTRACT_HEADINGS:
                flush_current()
                in_abstract = True
                in_references = False
                continue

            if in_references:
                reference_lines.append(text)
                continue

            flush_current()
            current = _new_section(int(block.get("level") or 1), clean)
            in_abstract = False
            continue

        if in_references:
            reference_lines.append(text)
            continue

        if block_type == "paragraph" and _looks_like_reference_entry_start(text):
            flush_current()
            in_references = True
            in_abstract = False
            reference_lines.append(text)
            continue

        if block_type == "table":
            current["tables"].append(block.get("id"))
            continue

        if lower.startswith("abstract:"):
            metadata["abstract"] = text.split(":", 1)[1].strip()
            continue

        keyword_match = KEYWORD_PATTERN.match(text)
        if keyword_match:
            metadata["keywords"] = [part.strip() for part in re.split(r"[;,]", keyword_match.group(1)) if part.strip()]
            continue

        if in_abstract:
            abstract_lines.append(text)
            continue

        if current is None:
            # Title pages, author blocks, copyright pages, and other preamble
            # should not become a synthetic body section.
            continue

        if lower not in FRONTMATTER_HEADINGS:
            current["paragraphs"].append(text)
            current.setdefault("paragraph_formats", []).append(block.get("source_format") or {})

    flush_current()

    if abstract_lines:
        metadata["abstract"] = " ".join(abstract_lines).strip()

    if reference_lines:
        references = _split_reference_lines(reference_lines)

    sections = [
        section
        for section in sections
        if section["paragraphs"] and section["title"].lower() not in FRONTMATTER_HEADINGS
    ]

    ir = {
        "metadata": metadata,
        "sections": sections or [_new_section(1, "Introduction")],
        "figures": [],
        "tables": [block for block in blocks if block.get("type") == "table"],
        "references": references,
        "footnotes": [],
        "source": {"path": str(source_path), "suffix": source_path.suffix.lower()},
    }
    _attach_semantic_ir_metadata(ir)
    return ir


def _fallback_text_extraction(source_path: Path) -> str:
    suffix = source_path.suffix.lower()
    if suffix == ".pdf":
        pdftotext = shutil_which("pdftotext")
        if pdftotext:
            completed = subprocess.run(
                [pdftotext, str(source_path), "-"],
                capture_output=True,
                text=True,
                check=False,
            )
            if completed.returncode == 0 and completed.stdout.strip():
                return completed.stdout

        with pdfplumber.open(source_path) as pdf:
            pages = [(page.extract_text() or "").strip() for page in pdf.pages]
        return "\n\n".join(page for page in pages if page)

    if suffix == ".docx":
        try:
            from docx import Document as WordDocument  # type: ignore
        except Exception as exc:
            raise RuntimeError("DOCX fallback requires python-docx") from exc
        document = WordDocument(str(source_path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n\n".join(paragraphs)

    return source_path.read_text(encoding="utf-8", errors="replace")


def _build_ir_from_markdown(markdown: str, source_path: Path) -> Dict[str, Any]:
    lines = [line.rstrip() for line in markdown.splitlines()]
    title = _first_nonempty_line(lines)
    authors, affiliation, abstract, keywords = _extract_frontmatter(lines)
    references = _extract_references(lines)
    footnotes = _extract_footnotes(lines)
    sections = _extract_sections(lines, title)

    return {
        "metadata": {
            "title": title,
            "authors": authors,
            "affiliation": affiliation,
            "abstract": abstract,
            "keywords": keywords,
            "date": None,
        },
        "sections": sections,
        "figures": [],
        "tables": [],
        "references": references,
        "footnotes": footnotes,
        "source": {"path": str(source_path), "suffix": source_path.suffix.lower()},
    }


def _markdown_to_blocks(markdown: str) -> List[Dict[str, Any]]:
    blocks: List[Dict[str, Any]] = []
    paragraph_buffer: List[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        text = " ".join(part.strip() for part in paragraph_buffer if part.strip()).strip()
        if text:
            blocks.append({"type": "paragraph", "text": text})
        paragraph_buffer = []

    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue
        if line.startswith("#"):
            flush_paragraph()
            level = min(max(len(line) - len(line.lstrip("#")), 1), 6)
            blocks.append({"type": "heading", "level": level, "text": line.lstrip("#").strip()})
            continue
        if _looks_like_heading_text(line):
            flush_paragraph()
            blocks.append({"type": "heading", "level": _heading_level_from_text(line), "text": _clean_heading_text(line)})
            continue
        paragraph_buffer.append(line)

    flush_paragraph()
    return blocks


def _ir_metrics(ir: Dict[str, Any], elapsed_ms: float) -> Dict[str, Any]:
    sections = ir.get("sections") or []
    paragraphs = [paragraph for section in sections for paragraph in section.get("paragraphs", [])]
    section_titles = [str(section.get("title") or "") for section in sections]
    return {
        "elapsed_ms": elapsed_ms,
        "section_count": len(sections),
        "paragraph_count": len(paragraphs),
        "reference_count": len(ir.get("references") or []),
        "table_count": len(ir.get("tables") or []),
        "has_abstract": bool((ir.get("metadata") or {}).get("abstract")),
        "avg_paragraph_words": round(sum(len(str(p).split()) for p in paragraphs) / max(1, len(paragraphs)), 2),
        "empty_title_count": sum(1 for title in section_titles if not title.strip()),
        "section_titles": section_titles[:20],
    }


def _ir_metric_deltas(
    legacy_ir: Dict[str, Any],
    chunker_ir: Dict[str, Any],
    legacy_ms: float,
    chunker_ms: float,
) -> Dict[str, Any]:
    legacy_metrics = _ir_metrics(legacy_ir, legacy_ms)
    chunker_metrics = _ir_metrics(chunker_ir, chunker_ms)
    numeric_keys = [
        "elapsed_ms",
        "section_count",
        "paragraph_count",
        "reference_count",
        "table_count",
        "avg_paragraph_words",
        "empty_title_count",
    ]
    return {
        key: round(chunker_metrics[key] - legacy_metrics[key], 2)
        for key in numeric_keys
        if isinstance(legacy_metrics.get(key), (int, float)) and isinstance(chunker_metrics.get(key), (int, float))
    }


def _ir_quality_score(ir: Dict[str, Any]) -> float:
    sections = ir.get("sections") or []
    paragraphs = [paragraph for section in sections for paragraph in section.get("paragraphs", []) or []]
    references = ir.get("references") or []
    metadata = ir.get("metadata") or {}

    section_count = len(sections)
    paragraph_count = len(paragraphs)
    avg_words = sum(len(str(paragraph).split()) for paragraph in paragraphs) / max(1, paragraph_count)
    single_paragraph_sections = sum(1 for section in sections if len(section.get("paragraphs") or []) <= 1)
    single_ratio = single_paragraph_sections / max(1, section_count)
    title_count = len({str(section.get("title") or "").strip().lower() for section in sections if str(section.get("title") or "").strip()})

    score = 0.0
    score += min(section_count, 18) * 2.0
    score += min(paragraph_count, 120) * 0.6
    score += min(len(references), 80) * 1.4
    score += 12.0 if metadata.get("abstract") else 0.0
    score += min(title_count, 18) * 1.0
    score += min(avg_words, 90) * 0.15
    if section_count > 25 and single_ratio > 0.55:
        score -= (section_count - 25) * 1.5
        score -= single_ratio * 20.0
    if paragraph_count <= 2:
        score -= 25.0
    if not sections:
        score -= 50.0
    return score


def _extract_frontmatter(lines: List[str]) -> Tuple[List[str], str | None, str | None, List[str]]:
    authors: List[str] = []
    affiliation: str | None = None
    abstract_lines: List[str] = []
    keywords: List[str] = []

    in_abstract = False
    for line in lines[:80]:
        stripped = line.strip()
        if not stripped:
            if in_abstract and abstract_lines:
                break
            continue

        if stripped.startswith("#"):
            heading = stripped.lstrip("#").strip().lower()
            if heading in ABSTRACT_HEADINGS:
                in_abstract = True
                continue
            if in_abstract:
                break

        keyword_match = KEYWORD_PATTERN.match(stripped)
        if keyword_match:
            keywords = [part.strip() for part in re.split(r"[;,]", keyword_match.group(1)) if part.strip()]
            continue

        if in_abstract:
            abstract_lines.append(stripped)
            continue

        if not authors and len(stripped.split()) <= 12 and "," in stripped:
            authors = [part.strip() for part in stripped.split(",") if part.strip()]
            continue

        if authors and affiliation is None and len(stripped.split()) <= 20 and "@" not in stripped:
            affiliation = stripped

    abstract = " ".join(abstract_lines).strip() or None
    return authors, affiliation, abstract, keywords


def _extract_sections(lines: List[str], title: str | None) -> List[Dict[str, Any]]:
    sections: List[Dict[str, Any]] = []
    current = _new_section(1, title or "Introduction")

    def flush() -> None:
        nonlocal current
        if current["title"] or current["paragraphs"] or current["figures"] or current["tables"]:
            current["paragraphs"] = [paragraph for paragraph in current["paragraphs"] if paragraph]
            sections.append(current)

    paragraph_buffer: List[str] = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            if paragraph_buffer:
                current["paragraphs"].append(" ".join(paragraph_buffer).strip())
                paragraph_buffer = []
            continue

        if line.startswith("#"):
            if paragraph_buffer:
                current["paragraphs"].append(" ".join(paragraph_buffer).strip())
                paragraph_buffer = []
            flush()
            level = min(max(len(line) - len(line.lstrip("#")), 1), 3)
            heading = line.lstrip("#").strip()
            current = _new_section(level, heading)
            continue

        lower = line.lower()
        if lower in REFERENCE_HEADINGS:
            break
        if lower.startswith("figure "):
            paragraph_buffer.append(line)
            continue
        paragraph_buffer.append(line)

    if paragraph_buffer:
        current["paragraphs"].append(" ".join(paragraph_buffer).strip())
    flush()

    return [section for section in sections if section["title"] or section["paragraphs"]]


def _extract_references(lines: List[str]) -> List[str]:
    references: List[str] = []
    in_references = False
    buffer: List[str] = []

    for raw_line in lines:
        line = raw_line.strip()
        lower = line.lower()
        if lower in REFERENCE_HEADINGS:
            in_references = True
            continue
        if not in_references:
            continue
        if not line:
            if buffer:
                references.append(" ".join(buffer).strip())
                buffer = []
            continue
        buffer.append(line)

    if buffer:
        references.append(" ".join(buffer).strip())

    return references


def _extract_footnotes(lines: List[str]) -> List[str]:
    return [line.strip() for line in lines if re.match(r"^\[\^.+\]:", line.strip())]


def _first_block_text(blocks: List[Dict[str, Any]]) -> str | None:
    title_candidates = [
        block
        for block in blocks[:30]
        if block.get("type") in {"title", "heading"}
        and str(block.get("text") or "").strip()
        and not _looks_like_metadata_or_sidebar_line(str(block.get("text") or ""))
    ]
    if title_candidates:
        largest = max(title_candidates, key=lambda block: float((block.get("source_format") or {}).get("font_size_pt") or 0))
        largest_size = float((largest.get("source_format") or {}).get("font_size_pt") or 0)
        if largest_size >= 12 or largest.get("type") == "title":
            matching = [
                block
                for block in title_candidates
                if abs(float((block.get("source_format") or {}).get("font_size_pt") or 0) - largest_size) <= 0.5
            ][:3]
            return " ".join(_clean_heading_text(str(block.get("text") or "")) for block in matching).strip()

    title_parts: List[str] = []
    for block in blocks[:12]:
        text = str(block.get("text") or "").strip()
        if not text:
            continue
        lower = text.lower()
        if lower in ABSTRACT_HEADINGS or lower in FRONTMATTER_HEADINGS:
            break
        if _looks_like_running_header(text) or _is_page_number(text) or _looks_like_metadata_or_sidebar_line(text):
            continue
        if len(text.split()) <= 14 and (text.isupper() or block.get("type") in {"title", "heading"}):
            title_parts.append(_clean_heading_text(text))
            continue
        if title_parts:
            break

    if title_parts:
        return " ".join(title_parts).strip()

    for block in blocks[:30]:
        text = str(block.get("text") or "").strip()
        if text and not _looks_like_running_header(text) and not _looks_like_metadata_or_sidebar_line(text) and block.get("type") in {"title", "heading", "paragraph"} and len(text.split()) <= 30:
            return _clean_heading_text(text)
    for block in blocks[:30]:
        text = str(block.get("text") or "").strip()
        if text:
            return text
    return None


def _looks_like_running_header(text: str) -> bool:
    return bool(re.search(r"\s+\d{1,4}$", text.strip())) and len(text.split()) <= 10


def _line_fingerprint(text: str) -> str:
    value = re.sub(r"\d+", "#", text.lower())
    value = re.sub(r"[^a-z# ]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def _is_page_number(text: str) -> bool:
    stripped = text.strip()
    return bool(re.fullmatch(r"(?:page\s*)?\d{1,4}", stripped, re.IGNORECASE))


def _is_frontmatter_page(lines: List[Dict[str, Any]]) -> bool:
    page_text = " ".join(line["text"] for line in lines[:18]).lower()
    if any(heading in page_text for heading in FRONTMATTER_HEADINGS):
        return True
    toc_markers = sum(1 for line in lines if re.search(r"_{4,}\s*\d+\s*$|\.{4,}\s*\d+\s*$", line["text"]))
    return toc_markers >= 4


def _is_pdf_heading(line: Dict[str, Any], body_size: float) -> bool:
    text = line["text"].strip()
    words = text.split()
    if _looks_like_metadata_or_sidebar_line(text):
        return False
    if len(words) > 16 or len(text) > 160:
        return False
    if line.get("size", 0) and line.get("size", 0) < body_size * 0.9 and _clean_heading_text(text).lower() not in COMMON_SECTION_HEADINGS:
        return False
    if re.match(r"^(?:[•\-*o]\s+|[a-z]\))", text, re.IGNORECASE):
        return False
    if ":" in text and len(words) > 8:
        return False
    clean = _clean_heading_text(text).lower()
    if clean in COMMON_SECTION_HEADINGS:
        return True
    if re.match(r"^(?:[ivxlcdm]+\.|[a-z]\.|chapter\s+\d+|\d+(?:\.\d+)*\.?)\s+\S+", text, re.IGNORECASE):
        return True
    if line.get("size", 0) >= body_size * 1.15 and _heading_shape(text):
        return True
    if line.get("bold") and _heading_shape(text):
        return True
    return False


def _looks_like_heading_text(text: str) -> bool:
    stripped = text.strip()
    if _looks_like_metadata_or_sidebar_line(stripped):
        return False
    if not stripped or len(stripped) > 160 or len(stripped.split()) > 16:
        return False
    clean = _clean_heading_text(stripped).lower()
    if clean in COMMON_SECTION_HEADINGS:
        return True
    return bool(re.match(r"^(?:[ivxlcdm]+\.|[a-z]\.|chapter\s+\d+|\d+(?:\.\d+)*\.?)\s+\S+", stripped, re.IGNORECASE))


def _looks_like_metadata_or_sidebar_line(text: str) -> bool:
    value = re.sub(r"\s+", " ", str(text or "").strip())
    lower = value.lower()
    if not value:
        return False
    if any(token in lower for token in ("doi:", "http://", "https://", "www.", "@", "copyright", "creative commons")):
        return True
    if re.match(r"^(received|revised|accepted|published)\s*:", lower):
        return True
    if re.match(r"^(how to cite this paper|article|correspondence|\*correspondence)\b", lower):
        return True
    if re.search(r"\b(department|university|institute|school|college|usa|inc)\b", lower) and len(value.split()) <= 14:
        return True
    if re.search(r"\bjournal\b", lower) and re.search(r"\b20\d{2}\b", lower):
        return True
    if re.fullmatch(r"[,.\d\s*]+", value):
        return True
    return False


def _should_skip_pdf_body_line(line: Dict[str, Any], body_size: float) -> bool:
    text = str(line.get("text") or "").strip()
    if line.get("zone") in {"header", "footer", "sidebar", "marginalia"}:
        return True
    if _looks_like_metadata_or_sidebar_line(text):
        return True
    if line.get("size", 0) and line.get("size", 0) < body_size * 0.82:
        return True
    return False


def _heading_shape(text: str) -> bool:
    stripped = text.strip()
    words = stripped.split()
    if not words or len(words) > 12:
        return False
    if stripped[-1:] in {".", ",", ";"}:
        return False
    alpha_words = [word for word in words if re.search(r"[A-Za-z]", word)]
    if not alpha_words:
        return False
    capitalized = sum(1 for word in alpha_words if word[:1].isupper() or word.isupper())
    return capitalized / len(alpha_words) >= 0.45


def _heading_level_from_text(text: str) -> int:
    stripped = text.strip()
    if re.match(r"^(?:chapter\s+\d+|[ivxlcdm]+\.)\s+", stripped, re.IGNORECASE):
        return 1
    if re.match(r"^\d+\.\d+\.\d+\s+", stripped):
        return 3
    if re.match(r"^(?:\d+\.\d+|[a-z]\.)\s+", stripped, re.IGNORECASE):
        return 2
    return 1


def _clean_heading_text(text: str) -> str:
    value = re.sub(r"^\s*(?:[ivxlcdm]+\.|[a-z]\.|chapter\s+\d+\s*:|\d+(?:\.\d+)*\.?)\s*", "", text.strip(), flags=re.IGNORECASE)
    value = re.sub(r"\s+", " ", value).strip(" -_\t")
    return value or text.strip()


def _split_reference_lines(lines: List[str]) -> List[str]:
    text = "\n".join(line.strip() for line in lines if line.strip())
    if not text:
        return []
    if re.search(r"\[\d+\]", text):
        parts = re.split(r"(?=\s*\[\d+\])", text)
        return [" ".join(part.split()) for part in parts if part.strip()]

    entries: List[str] = []
    current: List[str] = []
    entry_start = re.compile(r"^[A-Z][A-Za-z'`-]+,\s+")
    for line in text.splitlines():
        stripped = line.strip()
        if current and entry_start.match(stripped) and re.search(r"[.)\]]$", current[-1]):
            entries.append(" ".join(current))
            current = [stripped]
        else:
            current.append(stripped)
    if current:
        entries.append(" ".join(current))
    return entries


def _looks_like_reference_entry_start(text: str) -> bool:
    return bool(re.match(r"^\s*\[\d+\]\s+\S+", str(text or "")))


def _extract_figures_with_pymupdf(source_path: Path, image_dir: Path, warnings: List[str]) -> List[Dict[str, Any]]:
    try:
        import pymupdf as fitz
    except Exception:
        try:
            import fitz  # type: ignore
        except Exception as exc:
            warnings.append(f"PyMuPDF not available for figure extraction: {exc}")
            return []

    figures: List[Dict[str, Any]] = []
    try:
        document = fitz.open(str(source_path))
        for page_index, page in enumerate(document):
            for image_index, image in enumerate(page.get_images(full=True), start=1):
                xref = image[0]
                pixmap = fitz.Pixmap(document, xref)
                if pixmap.n - pixmap.alpha > 3:
                    pixmap = fitz.Pixmap(fitz.csRGB, pixmap)
                file_path = image_dir / f"fig_{page_index + 1}_{xref}_{image_index}.png"
                pixmap.save(str(file_path))
                figures.append(
                    {
                        "id": f"fig-{page_index + 1}-{image_index}",
                        "page": page_index + 1,
                        "caption": None,
                        "image_path": str(file_path),
                    }
                )
    except Exception as exc:
        warnings.append(f"Figure extraction failed: {exc}")
    return figures


def _extract_tables_with_pdfplumber(source_path: Path, warnings: List[str]) -> List[Dict[str, Any]]:
    tables: List[Dict[str, Any]] = []
    try:
        with pdfplumber.open(source_path) as pdf:
            table_id = 1
            for page_index, page in enumerate(pdf.pages, start=1):
                extracted_tables = page.extract_tables() or []
                for table in extracted_tables:
                    rows = [[cell or "" for cell in row] for row in table if row]
                    if not rows:
                        continue
                    tables.append(
                        {
                            "id": f"table-{table_id}",
                            "page": page_index,
                            "caption": None,
                            "rows": rows,
                        }
                    )
                    table_id += 1
    except Exception as exc:
        warnings.append(f"Table extraction fallback failed: {exc}")
    return tables


def _extract_pdf_references_from_text(source_path: Path, warnings: List[str]) -> List[str]:
    try:
        with pdfplumber.open(source_path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception as exc:
        warnings.append(f"Raw PDF reference recovery failed: {exc}")
        return []

    match = re.search(r"(?im)^\s*References\s*$", text)
    if not match:
        return []
    tail = text[match.end():]
    tail = re.sub(r"(?m)^\s*\S.+?\s+\d+\s+of\s+\d+\s*$", "", tail)
    tail = re.sub(r"(?m)^\s*(?:Universal Journal.+|www\.scipublications\.com.+|DOI:.+)\s*$", "", tail)
    return _split_reference_lines(tail.splitlines())


def _new_section(level: int, title: str) -> Dict[str, Any]:
    return {
        "level": level,
        "title": title,
        "paragraphs": [],
        "paragraph_formats": [],
        "figures": [],
        "tables": [],
    }


def _first_nonempty_line(lines: List[str]) -> str | None:
    for line in lines:
        stripped = line.strip().lstrip("#").strip()
        if stripped:
            return stripped
    return None


def _nougat_command() -> List[str] | None:
    executable = str(get_env("NOUGAT_COMMAND") or "").strip()
    if executable:
        return [executable]
    if shutil_which("nougat"):
        return ["nougat"]
    try:
        import importlib.util

        if importlib.util.find_spec("nougat") is not None:
            return [sys.executable, "-m", "nougat"]
    except Exception:
        return None
    return None


def _trim_warning(value: str, max_chars: int = 500) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rstrip() + "..."


def _int_env(name: str, default: int) -> int:
    raw = get_env(name)
    if raw is None:
        return default
    try:
        value = int(str(raw).strip())
        return value if value > 0 else default
    except Exception:
        return default


def _truthy_env(name: str) -> bool:
    raw = get_env(name)
    return str(raw or "").strip().lower() in {"1", "true", "yes", "on"}


def shutil_which(command: str) -> str | None:
    import shutil

    return shutil.which(command)
